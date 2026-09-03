import io
import csv
import json
import hashlib
import hmac
import re
import shutil
import subprocess
import urllib.parse
import urllib.request
from datetime import datetime, timedelta, timezone
from pathlib import Path
from uuid import uuid4

from fastapi import APIRouter, Depends, File, Form, HTTPException, Query, UploadFile, status
from fastapi.responses import FileResponse, Response, StreamingResponse
from sqlalchemy import case, func
from sqlalchemy.exc import SQLAlchemyError
from sqlalchemy.orm import Session

from app.api.deps import get_current_reviewer, get_current_user, get_db
from app.core.config import settings
from app.models.paper import Paper
from app.models.paper_workflow import PaperVersion, PaperReviewLog, PaperWorkflowEvent
from app.models.department import Department
from app.models.department_supervisor import DepartmentSupervisor
from app.models.user import User
from app.models.user_role import UserRole
from app.schemas.paper import PaperCreate, PaperRead, PaperReview, PaperStats, SupervisorReviewSummary
from app.services.notification_service import create_notification
from app.services.email_service import send_notification_email
from app.services.paper_service import (
    create_paper,
    get_paper,
    get_paper_stats,
    increment_download,
    increment_view,
    list_papers,
    review_paper,
)
from app.services.user_service import list_users, has_role, get_user_by_email, get_user_roles
from app.services.grading_service import classify_degree_level

router = APIRouter()
UPLOADS_DIR = Path(__file__).resolve().parents[3] / "uploads" / "papers"
UPLOADS_DIR.mkdir(parents=True, exist_ok=True)
EDITOR_TOKEN_TTL_SECONDS = 60 * 60


def _build_editor_token(*, paper_id: int, action: str, expires_in: int = EDITOR_TOKEN_TTL_SECONDS) -> str:
    exp = int(datetime.now(timezone.utc).timestamp()) + max(60, expires_in)
    payload = f"{paper_id}:{action}:{exp}"
    signature = hmac.new(
        settings.secret_key.encode("utf-8"),
        payload.encode("utf-8"),
        hashlib.sha256,
    ).hexdigest()
    return f"{payload}:{signature}"


def _verify_editor_token(*, token: str, paper_id: int, action: str) -> bool:
    try:
        token_paper_id, token_action, token_exp, token_sig = token.split(":", 3)
    except ValueError:
        return False
    if token_paper_id != str(paper_id) or token_action != action:
        return False
    try:
        exp = int(token_exp)
    except ValueError:
        return False
    now_ts = int(datetime.now(timezone.utc).timestamp())
    if exp < now_ts:
        return False
    raw_payload = f"{token_paper_id}:{token_action}:{token_exp}"
    expected_sig = hmac.new(
        settings.secret_key.encode("utf-8"),
        raw_payload.encode("utf-8"),
        hashlib.sha256,
    ).hexdigest()
    return hmac.compare_digest(token_sig, expected_sig)


def _sanitize_examiner_feedback_for_student(text: str | None) -> str | None:
    if not text:
        return None
    
    lines = text.split("\n")
    cleaned_lines = []
    
    for line in lines:
        l_str = line.strip().replace("\r", "")
        if not l_str:
            continue
        
        # Strip examiner headers and brackets
        l_clean = re.sub(r'\[[^\]]+\]\s*:?', '', l_str)
        # Strip any score, mark, or grade patterns
        l_clean = re.sub(r'-?\s*Final\s+Score:\s*[^•\n]+', '', l_clean, flags=re.IGNORECASE)
        l_clean = re.sub(r'-?\s*Score:\s*[^•\n]+', '', l_clean, flags=re.IGNORECASE)
        l_clean = re.sub(r'-?\s*Mark:\s*[^•\n]+', '', l_clean, flags=re.IGNORECASE)
        l_clean = re.sub(r'\([A-Za-z\s+]+\)', '', l_clean)
        
        l_clean = l_clean.strip().lstrip("-").rstrip("-").rstrip(":").lstrip(":").strip()
        if l_clean and l_clean not in {")", "(", "-", ":"}:
            cleaned_lines.append(l_clean)
            
    dedup = []
    seen = set()
    for l in cleaned_lines:
        if l not in seen:
            seen.add(l)
            dedup.append(l)
            
    if not dedup:
        return None
    return "[Examiner Evaluation & Required Corrections]\n" + "\n".join(dedup)


def _clean_examiner_corrections(text: str | None) -> str | None:
    if not text:
        return text
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    seen = set()
    deduped = []
    for line in lines:
        if line not in seen:
            seen.add(line)
            deduped.append(line)
    return "\n\n".join(deduped)


def _to_paper_read(paper: Paper, db: Session | None = None, current_user: User | None = None) -> PaperRead:
    # Check if user can view examiner marks/scores
    can_view_marks = False
    if current_user:
        if getattr(current_user, "is_admin", False) or (db and (has_role(db, current_user, "system_admin") or has_role(db, current_user, "dean") or has_role(db, current_user, "head_library"))):
            can_view_marks = True
        elif paper.internal_examiner_id == current_user.id or paper.external_examiner_id == current_user.id or paper.supervisor_id == current_user.id:
            can_view_marks = True
        elif db:
            user_dept = (current_user.department or "").strip().lower()
            paper_dept = ""
            if paper.department:
                paper_dept = (paper.department.name or "").strip().lower()
            elif paper.discipline:
                paper_dept = (paper.discipline or "").strip().lower()
            
            is_hod_or_coord = has_role(db, current_user, "hod") or has_role(db, current_user, "project_coordinator")
            if is_hod_or_coord and user_dept and paper_dept and user_dept == paper_dept:
                can_view_marks = True

    internal_score = paper.internal_score if can_view_marks else None
    external_score = paper.external_score if can_view_marks else None
    examiner_result_file_name = paper.examiner_result_file_name if can_view_marks else None
    internal_result_file_name = paper.internal_result_file_name if can_view_marks else None
    external_result_file_name = paper.external_result_file_name if can_view_marks else None

    # Classify degree level
    stu_user = db.query(User).filter(User.id == paper.created_by_id).first() if (db and paper.created_by_id) else None
    degree_level = classify_degree_level(paper=paper, student_user=stu_user, db=db)

    # Sanitize examiner corrections for student
    raw_corrections = paper.examiner_corrections
    if not can_view_marks and raw_corrections:
        examiner_corrections = _sanitize_examiner_feedback_for_student(raw_corrections)
    else:
        examiner_corrections = raw_corrections

    steps_list = []
    from sqlalchemy.orm import object_session
    session = db or object_session(paper)
    if session:
        from app.models.thesis_system import Step
        raw_steps = session.query(Step).filter(Step.thesis_id == paper.id).order_by(Step.step_number.asc()).all()
        from app.schemas.paper import StepRead
        steps_list = [StepRead.model_validate(s) for s in raw_steps]

    return PaperRead(
        id=paper.id,
        title=paper.title,
        abstract=paper.abstract,
        status=paper.status,
        discipline=paper.discipline,
        university=paper.university,
        year=paper.year,
        document_type=paper.document_type,
        publication_type=paper.publication_type,
        doi=paper.doi,
        license=paper.license,
        file_name=paper.file_name,
        file_size=paper.file_size,
        mime_type=paper.mime_type,
        views=paper.views or 0,
        downloads=paper.downloads or 0,
        citations=paper.citations or 0,
        rating=paper.rating,
        review_comments=paper.review_comments,
        supervisor_id=paper.supervisor_id,
        department_id=paper.department_id,
        abstract_word_count=paper.abstract_word_count,
        work_mode=paper.work_mode or "individual",
        created_at=paper.created_at,
        authors=[a for a in paper.authors],
        tags=[pt.tag.name for pt in paper.tags if pt.tag],
        # 5-phase properties
        project_coordinator_id=paper.project_coordinator_id,
        internal_examiner_id=paper.internal_examiner_id,
        external_examiner_id=paper.external_examiner_id,
        ch1_student_done=paper.ch1_student_done,
        ch2_student_done=paper.ch2_student_done,
        ch3_student_done=paper.ch3_student_done,
        ch4_student_done=paper.ch4_student_done,
        ch5_student_done=paper.ch5_student_done,
        ch1_supervisor_approved=paper.ch1_supervisor_approved,
        ch2_supervisor_approved=paper.ch2_supervisor_approved,
        ch3_supervisor_approved=paper.ch3_supervisor_approved,
        ch4_supervisor_approved=paper.ch4_supervisor_approved,
        ch5_supervisor_approved=paper.ch5_supervisor_approved,
        combined_thesis_student_done=paper.combined_thesis_student_done,
        combined_thesis_supervisor_approved=paper.combined_thesis_supervisor_approved,
        internal_score=internal_score,
        external_score=external_score,
        examiner_corrections=examiner_corrections,
        examiner_result_file_name=examiner_result_file_name,
        internal_result_file_name=internal_result_file_name,
        external_result_file_name=external_result_file_name,
        lecturer_approved_at=paper.lecturer_approved_at,
        project_coordinator_approved_at=paper.project_coordinator_approved_at,
        hod_approved_at=paper.hod_approved_at,
        steps=steps_list,
        degree_level=degree_level,
    )


def _notify_roles(
    db: Session,
    *,
    roles: set[str],
    paper_id: int,
    message: str,
    department: str | None = None,
) -> None:
    dept = (department or "").strip().lower()
    recipients = [
        user
        for user in list_users(db, limit=500, is_active=True)
        if any(has_role(db, user, role) for role in roles)
        and (not dept or (user.department or "").strip().lower() == dept)
    ]
    # Fallback: if department-filtered search finds nobody (often due naming mismatch),
    # notify role users without department restriction so workflow never stalls silently.
    if dept and not recipients:
        recipients = [
            user
            for user in list_users(db, limit=500, is_active=True)
            if any(has_role(db, user, role) for role in roles)
        ]
    for recipient in recipients:
        create_notification(
            db,
            user_id=recipient.id,
            paper_id=paper_id,
            ntype="workflow_update",
            message=message,
        )


def _resolve_reviewer_role(db: Session, user: User) -> str:
    ordered_roles = ["librarian", "hod", "project_coordinator", "project_supervisor", "lecturer"]
    for role in ordered_roles:
        if has_role(db, user, role):
            return role
    return (user.role or "").strip().lower()


def _is_lecturer_submitter(db: Session, user: User) -> bool:
    return has_role(db, user, "lecturer")


def _can_submit_paper(db: Session, user: User) -> bool:
    return has_role(db, user, "student") or has_role(db, user, "member") or has_role(db, user, "lecturer")


def _doi_prefix() -> str:
    prefix = (settings.doi_prefix or "10.99999").strip().rstrip("/")
    return prefix or "10.99999"


def _slug_doi_part(value: str | None, fallback: str) -> str:
    raw = (value or "").strip().lower()
    if not raw:
        return fallback
    compact = re.sub(r"[^a-z0-9]+", "-", raw).strip("-")
    return compact or fallback


def _generate_paper_doi(paper: Paper) -> str:
    school = _slug_doi_part(paper.university, "school")
    department = _slug_doi_part(paper.discipline, "department")
    year = int(paper.year or datetime.now(timezone.utc).year)
    return f"{_doi_prefix()}/murrs.{school}.{department}.{year}.{paper.id}"


def _ensure_paper_doi(db: Session, paper: Paper) -> None:
    if paper.status != "approved":
        return
    if (paper.doi or "").strip():
        return
    paper.doi = _generate_paper_doi(paper)
    db.add(paper)


def _notify_student(db: Session, paper: Paper, message: str) -> None:
    if paper.created_by_id:
        create_notification(
            db,
            user_id=paper.created_by_id,
            paper_id=paper.id,
            ntype="workflow_update",
            message=message,
        )
        return

    # Fallback for imported/legacy records where created_by_id is missing:
    # send email directly to paper author emails so they still receive updates.
    seen: set[str] = set()
    for author in (paper.authors or []):
        email = ((author.email or "").strip().lower() if author else "")
        if not email or email in seen:
            continue
        seen.add(email)
        send_notification_email(
            to_email=email,
            to_name=(author.name if author else None),
            subject="MURRS workflow update",
            message=message,
        )


def _notify_hod_and_coordinators(db: Session, paper: Paper, phase_message: str) -> None:
    # Notify HODs and Project Coordinators of the paper's department
    dept_name = ""
    if paper.department:
        dept_name = paper.department.name
    elif paper.discipline:
        dept_name = paper.discipline
        
    _notify_roles(
        db,
        roles={"hod", "project_coordinator"},
        paper_id=paper.id,
        department=dept_name,
        message=phase_message,
    )


def _notify_assigned_lecturer(db: Session, paper: Paper, submitted_by: User) -> None:
    if paper.supervisor_id:
        create_notification(
            db,
            user_id=paper.supervisor_id,
            paper_id=paper.id,
            ntype="workflow_update",
            message=f"New submission assigned to you by {submitted_by.full_name or submitted_by.email}: {paper.title}",
        )
        return
    _notify_roles(
        db,
        roles={"lecturer"},
        paper_id=paper.id,
        department=submitted_by.department,
        message=f"New submission awaiting lecturer review: {paper.title}",
    )


def _notify_hod_and_project_coordinator_on_student_submission(
    db: Session,
    *,
    paper: Paper,
    submitted_by: User,
) -> None:
    _notify_roles(
        db,
        roles={"hod", "project_coordinator"},
        paper_id=paper.id,
        department=submitted_by.department,
        message=(
            "New student paper submission awaiting approval workflow.\n"
            f"Title: {paper.title}\n"
            f"Student: {submitted_by.full_name or submitted_by.email}\n"
            "Status: pending lecturer review"
        ),
    )


def _notify_supervisor_proposal_summary(db: Session, paper: Paper) -> None:
    """Send project abstract summary notification & email to student's supervisor when proposal is approved."""
    supervisor_ids: set[int] = set()
    if paper.supervisor_id:
        supervisor_ids.add(paper.supervisor_id)
    if hasattr(paper, "supervisors") and paper.supervisors:
        for s in paper.supervisors:
            if getattr(s, "user_id", None):
                supervisor_ids.add(s.user_id)

    if not supervisor_ids:
        return

    student_name = (
        paper.created_by.full_name or paper.created_by.email
        if paper.created_by
        else "Student"
    )
    abstract_text = (paper.abstract or "").strip() or "No abstract summary provided."
    summary_message = (
        f"Proposal Approved: Student {student_name}'s proposal '{paper.title}' has been approved.\n\n"
        f"Brief Project Summary (Abstract):\n{abstract_text}"
    )

    for s_id in supervisor_ids:
        create_notification(
            db,
            user_id=s_id,
            paper_id=paper.id,
            ntype="proposal_approved_summary",
            message=summary_message,
        )
        sup_user = db.query(User).filter(User.id == s_id).first()
        if sup_user and sup_user.email:
            send_notification_email(
                to_email=sup_user.email,
                to_name=sup_user.full_name or sup_user.email,
                subject=f"Approved Student Project Summary: {paper.title}",
                message=summary_message,
            )



def _dispatch_overdue_review_alerts(db: Session) -> None:
    now = datetime.now(timezone.utc)
    cutoff = now - timedelta(days=3)
    overdue_rows = (
        db.query(Paper, User)
        .join(User, Paper.created_by_id == User.id)
        .filter(
            Paper.status == "pending_lecturer",
            Paper.created_at <= cutoff,
            Paper.lecturer_overdue_alert_sent_at.is_(None),
        )
        .all()
    )

    if not overdue_rows:
        return

    for paper, owner in overdue_rows:
        # Only student/member submissions require this escalation.
        if not (has_role(db, owner, "student") or has_role(db, owner, "member")):
            paper.lecturer_overdue_alert_sent_at = now
            db.add(paper)
            continue

        _notify_roles(
            db,
            roles={"hod", "project_coordinator"},
            paper_id=paper.id,
            department=owner.department,
            message=(
                "Overdue lecturer review alert (3+ days).\n"
                f"Title: {paper.title}\n"
                f"Student: {owner.full_name or owner.email}\n"
                f"Submitted: {paper.created_at.isoformat() if paper.created_at else '-'}\n"
                "The assigned lecturer has not completed review."
            ),
        )
        paper.lecturer_overdue_alert_sent_at = now
        db.add(paper)

    db.commit()


def _is_annotation_participant(paper: Paper, user: User) -> bool:
    is_author = paper.created_by_id == user.id
    is_assigned_supervisor = paper.supervisor_id == user.id
    is_paper_supervisor = any(s.user_id == user.id for s in paper.supervisors)
    return bool(user.is_admin or is_author or is_assigned_supervisor or is_paper_supervisor)


def _safe_match_department_id(db: Session, department_name: str | None) -> int | None:
    name = (department_name or "").strip()
    if not name:
        return None
    try:
        row = (
            db.query(Department.id)
            .filter(func.lower(Department.name) == name.lower())
            .first()
        )
    except SQLAlchemyError:
        return None
    if not row:
        return None
    return int(row[0])


def _sha256_file(path: Path) -> str | None:
    try:
        return hashlib.sha256(path.read_bytes()).hexdigest()
    except Exception:
        return None


def _convert_document_to_pdf_or_fail(source_path: Path) -> Path:
    if source_path.suffix.lower() == ".pdf":
        return source_path
    if not source_path.exists() or not source_path.is_file():
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Source file for PDF conversion was not found")

    expected_output = UPLOADS_DIR / f"{source_path.stem}.pdf"
    soffice_binary = shutil.which("soffice") or shutil.which("soffice.exe")
    if not soffice_binary:
        # Fallback for Windows hosts where PATH is not propagated to the backend process.
        windows_candidates = [
            Path("C:/Program Files/LibreOffice/program/soffice.exe"),
            Path("C:/Program Files (x86)/LibreOffice/program/soffice.exe"),
        ]
        for candidate in windows_candidates:
            if candidate.exists():
                soffice_binary = str(candidate)
                break
    if not soffice_binary:
        # Fallback for common Linux install locations.
        linux_candidates = [Path("/usr/bin/soffice"), Path("/usr/local/bin/soffice")]
        for candidate in linux_candidates:
            if candidate.exists():
                soffice_binary = str(candidate)
                break

    try:
        result = subprocess.run(
            [
                soffice_binary or "soffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(UPLOADS_DIR),
                str(source_path),
            ],
            check=True,
            capture_output=True,
            text=True,
            timeout=180,
        )
    except FileNotFoundError as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="PDF conversion failed: LibreOffice (soffice) is not installed on the server.",
        ) from exc
    except subprocess.CalledProcessError as exc:
        detail = (exc.stderr or exc.stdout or "Unknown conversion error").strip()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"PDF conversion failed: {detail}",
        ) from exc
    except subprocess.TimeoutExpired as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="PDF conversion failed: conversion timed out.",
        ) from exc

    if not expected_output.exists() or not expected_output.is_file():
        stdout_text = (result.stdout or "").strip()
        stderr_text = (result.stderr or "").strip()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=(
                "PDF conversion failed: output file not generated. "
                f"stdout={stdout_text[:200]} stderr={stderr_text[:200]}"
            ),
        )

    unique_pdf = UPLOADS_DIR / f"{uuid4().hex}_{expected_output.name}"
    expected_output.replace(unique_pdf)
    return unique_pdf


def _next_paper_version_no(db: Session, paper_id: int) -> int:
    latest = (
        db.query(PaperVersion.version_no)
        .filter(PaperVersion.paper_id == paper_id)
        .order_by(PaperVersion.version_no.desc())
        .first()
    )
    return int(latest[0]) + 1 if latest else 1


def _record_paper_version(
    db: Session,
    *,
    paper: Paper,
    source: str,
    actor_id: int | None,
    note: str | None = None,
) -> None:
    if not paper.file_path:
        return
    file_path = Path(paper.file_path)
    file_hash = _sha256_file(file_path) if file_path.exists() else None
    version = PaperVersion(
        paper_id=paper.id,
        version_no=_next_paper_version_no(db, paper.id),
        source=source,
        file_name=paper.file_name or file_path.name,
        file_path=str(file_path),
        file_size=paper.file_size,
        mime_type=paper.mime_type,
        file_sha256=file_hash,
        note=note,
        uploaded_by_id=actor_id,
    )
    db.add(version)


def _record_review_log(
    db: Session,
    *,
    paper_id: int,
    reviewer_id: int | None,
    reviewer_role: str | None,
    decision: str,
    comments: str | None,
    from_status: str | None,
    to_status: str | None,
) -> None:
    db.add(
        PaperReviewLog(
            paper_id=paper_id,
            reviewer_id=reviewer_id,
            reviewer_role=reviewer_role,
            decision=decision,
            comments=comments,
            from_status=from_status,
            to_status=to_status,
        )
    )


def _record_workflow_event(
    db: Session,
    *,
    paper_id: int,
    event_type: str,
    actor_id: int | None = None,
    actor_role: str | None = None,
    from_status: str | None = None,
    to_status: str | None = None,
    message: str | None = None,
) -> None:
    db.add(
        PaperWorkflowEvent(
            paper_id=paper_id,
            event_type=event_type,
            from_status=from_status,
            to_status=to_status,
            actor_id=actor_id,
            actor_role=actor_role,
            message=message,
        )
    )


@router.get("/papers", response_model=list[PaperRead])
def read_papers(
    db: Session = Depends(get_db),
    q: str | None = None,
    discipline: str | None = None,
    university: str | None = None,
    year: int | None = None,
    publication_type: str | None = None,
    department_id: int | None = None,
    status_filter: str | None = Query(None, alias="status"),
    catalog: bool = Query(False),
    sort: str = Query("relevance"),
    skip: int = Query(0, ge=0),
    limit: int = Query(50, ge=1, le=200),
) -> list[PaperRead]:
    papers = list_papers(
        db,
        q=q,
        discipline=discipline,
        university=university,
        year=year,
        publication_type=publication_type,
        department_id=department_id,
        status=status_filter,
        sort=sort,
        skip=skip,
        limit=limit,
        catalog_mode=catalog,
    )
    return [_to_paper_read(p) for p in papers]


@router.get("/papers/stats", response_model=PaperStats)
def read_paper_stats(db: Session = Depends(get_db)) -> PaperStats:
    return PaperStats(**get_paper_stats(db))


@router.get("/papers/pipeline")
def get_pipeline_metrics(
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    """Returns Phase 1 to Phase 5 metrics and student lists for HOD, Coordinator, Dean, Admin."""
    reviewer_department = (current_user.department or "").strip().lower()

    query = db.query(Paper).join(User, Paper.created_by_id == User.id, isouter=True)

    # Departmental Isolation Rule: HOD and Coordinator only see their department
    is_hod_or_coord = has_role(db, current_user, "hod") or has_role(db, current_user, "project_coordinator")
    if is_hod_or_coord and not current_user.is_admin and reviewer_department:
        query = query.filter(func.lower(func.coalesce(User.department, "")) == reviewer_department)

    papers = query.all()

    phases = {
        "phase1_proposals": {"count": 0, "students": []},
        "phase2_allocation": {"count": 0, "students": []},
        "phase3_chapters": {"count": 0, "students": []},
        "phase4_examination": {"count": 0, "students": []},
        "phase5_signoff": {"count": 0, "students": []},
    }

    for p in papers:
        student_user = p.created_by
        student_id = student_user.school_id if student_user else f"GIMPA-ST-{p.id:03d}"
        student_name = student_user.full_name if student_user else (p.authors[0].name if p.authors else "Unknown Student")
        program = p.discipline or ((student_user.program or student_user.department) if student_user else "B.Sc. Computer Science")
        supervisor_name = p.supervisor.full_name if p.supervisor else "Unassigned"

        status = (p.status or "").lower()

        item = {
            "paper_id": p.id,
            "index_number": student_id or f"GIMPA-ST-{p.id:03d}",
            "student_name": student_name,
            "program": program or "Computer Science",
            "supervisor_name": supervisor_name,
            "title": p.title,
            "status": p.status,
        }

        if status in {"draft", "pending", "pending_hod", "phase1_proposal_submitted", "phase1_proposal_rejected"}:
            item["milestone_status"] = "Phase 1 — Proposal Submitted"
            phases["phase1_proposals"]["students"].append(item)
            phases["phase1_proposals"]["count"] += 1
        elif status in {"phase1_topic_accepted", "phase2_proposal_submitted", "phase2_proposal_accepted", "pending_coordinator", "pending_hod_and_coordinator", "phase2_pending_coordinator", "phase2_pending_supervisor"}:
            item["milestone_status"] = "Phase 2 — Proposal & Allocation"
            phases["phase2_allocation"]["students"].append(item)
            phases["phase2_allocation"]["count"] += 1
        elif status in {"pending_lecturer", "revision", "phase3_chapters", "phase3_steps_in_progress"}:
            item["milestone_status"] = "Phase 3 — Chapter Writing & Review"
            phases["phase3_chapters"]["students"].append(item)
            phases["phase3_chapters"]["count"] += 1
        elif status in {"pending_examiner", "phase4_pending_examiners", "phase4_marking"}:
            item["milestone_status"] = "Phase 4 — Examination & Marking"
            phases["phase4_examination"]["students"].append(item)
            phases["phase4_examination"]["count"] += 1
        elif status in {
            "approved_for_library",
            "approved",
            "phase5_corrections",
            "phase5_pending_supervisor",
            "phase5_pending_coordinator",
            "phase5_pending_hod",
            "phase5_pending_hod_and_coordinator",
            "phase5_approved_for_library",
            "phase5_published",
        }:
            item["milestone_status"] = "Phase 5 — Corrections & Final Sign-off"
            phases["phase5_signoff"]["students"].append(item)
            phases["phase5_signoff"]["count"] += 1
        else:
            item["milestone_status"] = f"In Progress ({p.status})"
            phases["phase3_chapters"]["students"].append(item)
            phases["phase3_chapters"]["count"] += 1

    return phases


@router.get("/papers/pending", response_model=list[PaperRead])
def read_pending_papers(
    db: Session = Depends(get_db),
    current_admin=Depends(get_current_reviewer),
) -> list[PaperRead]:
    _dispatch_overdue_review_alerts(db)
    reviewer_role = _resolve_reviewer_role(db, current_admin)
    reviewer_department = (current_admin.department or "").strip().lower()
    if reviewer_role in {"project_coordinator", "hod"} and not reviewer_department:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Reviewer department is not configured")
    # Universal examiner condition: Any paper where current_admin is assigned as examiner during Phase 4 marking
    assigned_as_examiner_cond = (
        (Paper.status == "phase4_marking")
        & (
            (Paper.internal_examiner_id == current_admin.id)
            | (Paper.external_examiner_id == current_admin.id)
        )
    )

    if reviewer_role in {"lecturer", "project_supervisor", "external_examiner"}:
        supervisor_cond = (
            Paper.status.in_({
                "pending_lecturer",
                "phase2_proposal_submitted",
                "phase3_chapters",
                "phase3_steps_in_progress",  # thesis visible after first step submitted
                "phase5_pending_supervisor",
            })
        )
        if reviewer_department:
            supervisor_cond = supervisor_cond & (
                (Paper.supervisor_id == current_admin.id)
                | (
                    (Paper.supervisor_id.is_(None))
                    & (func.lower(func.coalesce(User.department, "")) == reviewer_department)
                    & (Paper.created_by_id != current_admin.id)
                )
            )
        else:
            supervisor_cond = supervisor_cond & (Paper.supervisor_id == current_admin.id)
            
        examiner_cond = (
            (Paper.status == "phase4_marking")
            & (
                (Paper.internal_examiner_id == current_admin.id)
                | (Paper.external_examiner_id == current_admin.id)
            )
        )
        
        query = (
            db.query(Paper)
            .join(User, Paper.created_by_id == User.id, isouter=True)
            .filter(supervisor_cond | examiner_cond)
        )
        papers = query.order_by(Paper.created_at.desc(), Paper.id.desc()).limit(200).all()
    elif reviewer_role == "project_coordinator":
        papers = (
            db.query(Paper)
            .join(User, Paper.created_by_id == User.id, isouter=True)
            .filter(
                assigned_as_examiner_cond
                | (
                    (
                        ((Paper.status == "pending_hod_and_coordinator") & (Paper.project_coordinator_approved_at.is_(None)))
                        | (Paper.status == "pending_coordinator")
                        | (Paper.status == "phase1_proposal_submitted")
                        | (Paper.status == "phase2_pending_coordinator")
                        | (Paper.status == "phase2_pending_supervisor")
                        | (Paper.status == "phase4_pending_examiners")
                        | (Paper.status == "phase4_marking")
                        | ((Paper.status == "phase5_pending_hod_and_coordinator") & (Paper.project_coordinator_approved_at.is_(None)))
                        | (Paper.status == "phase5_pending_coordinator")
                    )
                    & (func.lower(func.coalesce(User.department, "")) == reviewer_department)
                )
            )
            .order_by(Paper.created_at.desc(), Paper.id.desc())
            .limit(200)
            .all()
        )
    elif reviewer_role == "hod":
        papers = (
            db.query(Paper)
            .join(User, Paper.created_by_id == User.id, isouter=True)
            .filter(
                assigned_as_examiner_cond
                | (
                    (
                        ((Paper.status == "pending_hod_and_coordinator") & (Paper.hod_approved_at.is_(None)))
                        | (Paper.status == "pending_hod")
                        | (Paper.status == "phase1_proposal_submitted")
                        | (Paper.status == "phase2_pending_coordinator")
                        | (Paper.status == "phase2_pending_supervisor")
                        | (Paper.status == "phase4_pending_examiners")
                        | (Paper.status == "phase4_marking")
                        | ((Paper.status == "phase5_pending_hod_and_coordinator") & (Paper.hod_approved_at.is_(None)))
                        | (Paper.status == "phase5_pending_hod")
                    )
                    & (func.lower(func.coalesce(User.department, "")) == reviewer_department)
                )
            )
            .order_by(Paper.created_at.desc(), Paper.id.desc())
            .limit(200)
            .all()
        )
    elif reviewer_role == "librarian":
        papers = list_papers(db, status="approved_for_library", sort="newest", limit=200)
    else:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Reviewer role is not allowed")
    return [_to_paper_read(p, db) for p in papers]


@router.get("/papers/reviewed", response_model=list[PaperRead])
def read_reviewed_papers(
    db: Session = Depends(get_db),
    current_admin=Depends(get_current_reviewer),
) -> list[PaperRead]:
    reviewer_role = _resolve_reviewer_role(db, current_admin)
    reviewer_department = (current_admin.department or "").strip().lower()
    if reviewer_role in {"project_coordinator", "hod"} and not reviewer_department:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Reviewer department is not configured")

    if reviewer_role in {"lecturer", "project_supervisor"}:
        papers = (
            db.query(Paper)
            .filter(Paper.lecturer_approved_by_id == current_admin.id)
            .order_by(Paper.created_at.desc(), Paper.id.desc())
            .limit(200)
            .all()
        )
    elif reviewer_role == "project_coordinator":
        papers = (
            db.query(Paper)
            .join(User, Paper.created_by_id == User.id, isouter=True)
            .filter(Paper.project_coordinator_approved_by_id == current_admin.id)
            .filter(func.lower(func.coalesce(User.department, "")) == reviewer_department)
            .order_by(Paper.created_at.desc(), Paper.id.desc())
            .limit(200)
            .all()
        )
    elif reviewer_role == "hod":
        papers = (
            db.query(Paper)
            .join(User, Paper.created_by_id == User.id, isouter=True)
            .filter(Paper.hod_approved_by_id == current_admin.id)
            .filter(func.lower(func.coalesce(User.department, "")) == reviewer_department)
            .order_by(Paper.created_at.desc(), Paper.id.desc())
            .limit(200)
            .all()
        )
    elif reviewer_role == "librarian":
        papers = (
            db.query(Paper)
            .filter(Paper.status == "approved", Paper.reviewed_by_id == current_admin.id)
            .order_by(Paper.created_at.desc(), Paper.id.desc())
            .limit(200)
            .all()
        )
    else:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Reviewer role is not allowed")

    return [_to_paper_read(p, db) for p in papers]


@router.get("/papers/department/supervisor-reviewed", response_model=list[PaperRead])
def read_department_supervisor_reviewed_papers(
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
) -> list[PaperRead]:
    allowed = (
        current_user.is_admin
        or has_role(db, current_user, "hod")
        or has_role(db, current_user, "project_coordinator")
        or has_role(db, current_user, "system_admin")
    )
    if not allowed:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Not allowed")

    reviewer_department = (current_user.department or "").strip().lower()
    if not reviewer_department:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Reviewer department is not configured")

    dept_ids = [
        int(row[0])
        for row in (
            db.query(Department.id)
            .filter(func.lower(func.coalesce(Department.name, "")) == reviewer_department)
            .all()
        )
    ]

    supervisor_ids: set[int] = set(
        int(row[0])
        for row in (
            db.query(DepartmentSupervisor.supervisor_user_id)
            .filter(
                DepartmentSupervisor.department_id.in_(dept_ids) if dept_ids else False,
                DepartmentSupervisor.active.is_(True),
            )
            .all()
        )
    )

    # Fallback for environments where department_supervisors mapping is incomplete.
    fallback_project_supervisors = (
        db.query(User.id)
        .join(UserRole, UserRole.user_id == User.id)
        .filter(
            func.lower(func.coalesce(User.department, "")) == reviewer_department,
            func.lower(UserRole.role) == "project_supervisor",
            User.is_active.is_(True),
        )
        .all()
    )
    supervisor_ids.update(int(row[0]) for row in fallback_project_supervisors)

    if not supervisor_ids:
        return []

    papers = (
        db.query(Paper)
        .join(User, Paper.created_by_id == User.id, isouter=True)
        .filter(
            Paper.lecturer_approved_by_id.in_(list(supervisor_ids)),
            func.lower(func.coalesce(User.department, "")) == reviewer_department,
        )
        .order_by(Paper.lecturer_approved_at.desc().nullslast(), Paper.created_at.desc(), Paper.id.desc())
        .limit(500)
        .all()
    )
    return [_to_paper_read(p, db) for p in papers]


@router.get("/papers/department/supervisor-review-summary", response_model=list[SupervisorReviewSummary])
def read_department_supervisor_review_summary(
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
) -> list[SupervisorReviewSummary]:
    allowed = (
        current_user.is_admin
        or has_role(db, current_user, "hod")
        or has_role(db, current_user, "project_coordinator")
        or has_role(db, current_user, "system_admin")
    )
    if not allowed:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Not allowed")

    reviewer_department = (current_user.department or "").strip().lower()
    if not reviewer_department:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Reviewer department is not configured")

    dept_ids = [
        int(row[0])
        for row in (
            db.query(Department.id)
            .filter(func.lower(func.coalesce(Department.name, "")) == reviewer_department)
            .all()
        )
    ]

    supervisor_ids: set[int] = set()

    # 1. Department Supervisors mapping table
    if dept_ids:
        dept_sup_rows = (
            db.query(DepartmentSupervisor.supervisor_user_id)
            .filter(
                DepartmentSupervisor.department_id.in_(dept_ids),
                DepartmentSupervisor.active.is_(True),
            )
            .all()
        )
        supervisor_ids.update(int(row[0]) for row in dept_sup_rows)

    # 2. Users with supervisor/lecturer roles
    query_users = db.query(User.id).filter(User.is_active.is_(True))
    if not current_user.is_admin and reviewer_department:
        query_users = query_users.filter(func.lower(func.coalesce(User.department, "")) == reviewer_department)
    
    query_users = query_users.filter(
        func.lower(User.role).in_(["project_supervisor", "lecturer", "hod", "project_coordinator"])
    )
    supervisor_ids.update(int(row[0]) for row in query_users.all())

    # 3. Any user assigned as supervisor_id on papers in the department
    paper_sups = db.query(Paper.supervisor_id).filter(Paper.supervisor_id.isnot(None))
    if not current_user.is_admin and reviewer_department:
        paper_sups = paper_sups.join(User, Paper.created_by_id == User.id).filter(
            func.lower(func.coalesce(User.department, "")) == reviewer_department
        )
    supervisor_ids.update(int(row[0]) for row in paper_sups.all())

    if not supervisor_ids:
        return []

    supervisors = (
        db.query(User)
        .filter(User.id.in_(list(supervisor_ids)))
        .order_by(User.full_name.asc().nullslast(), User.email.asc())
        .all()
    )

    review_rows = (
        db.query(
            PaperReviewLog.reviewer_id,
            func.count(PaperReviewLog.id).label("reviews_done"),
            func.sum(case((PaperReviewLog.decision == "approve", 1), else_=0)).label("approvals_done"),
        )
        .join(Paper, Paper.id == PaperReviewLog.paper_id)
        .join(User, User.id == Paper.created_by_id)
        .filter(
            PaperReviewLog.reviewer_id.in_(list(supervisor_ids)),
            func.lower(func.coalesce(User.role, "")).in_(["student", "member"]),
        )
        .group_by(PaperReviewLog.reviewer_id)
        .all()
    )

    counts_by_supervisor: dict[int, tuple[int, int]] = {
        int(row[0]): (int(row[1] or 0), int(row[2] or 0)) for row in review_rows
    }

    # Also count assigned papers per supervisor for accurate reviews & approvals count
    assigned_paper_counts = (
        db.query(
            Paper.supervisor_id,
            func.count(Paper.id).label("assigned_total"),
            func.sum(
                case(
                    (
                        Paper.status.in_(
                            [
                                "phase5_pending_supervisor",
                                "phase5_pending_coordinator",
                                "phase5_pending_hod",
                                "phase5_pending_hod_and_coordinator",
                                "phase5_approved_for_library",
                                "approved",
                                "phase5_published",
                            ]
                        ),
                        1,
                    ),
                    else_=0,
                )
            ).label("approved_total"),
        )
        .filter(Paper.supervisor_id.in_(list(supervisor_ids)))
        .group_by(Paper.supervisor_id)
        .all()
    )
    assigned_dict = {
        int(row[0]): (int(row[1] or 0), int(row[2] or 0)) for row in assigned_paper_counts if row[0] is not None
    }

    summary: list[SupervisorReviewSummary] = []
    for supervisor in supervisors:
        log_reviews, log_approvals = counts_by_supervisor.get(supervisor.id, (0, 0))
        ass_total, ass_approved = assigned_dict.get(supervisor.id, (0, 0))
        
        total_reviews = max(log_reviews, ass_total)
        total_approvals = max(log_approvals, ass_approved)

        summary.append(
            SupervisorReviewSummary(
                supervisor_user_id=supervisor.id,
                supervisor_name=supervisor.full_name,
                supervisor_email=supervisor.email,
                department=supervisor.department,
                reviews_done=total_reviews,
                approvals_done=total_approvals,
                students_count=ass_total,
            )
        )

    summary.sort(key=lambda item: (-item.reviews_done, -item.approvals_done, item.supervisor_email))
    return summary


@router.get("/papers/revisions", response_model=list[PaperRead])
def read_revision_papers(
    db: Session = Depends(get_db),
    current_admin=Depends(get_current_reviewer),
) -> list[PaperRead]:
    reviewer_role = _resolve_reviewer_role(db, current_admin)
    reviewer_department = (current_admin.department or "").strip().lower()
    if reviewer_role in {"project_coordinator", "hod"} and not reviewer_department:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Reviewer department is not configured")

    if reviewer_role in {"lecturer", "project_supervisor", "librarian"}:
        papers = (
            db.query(Paper)
            .filter(
                Paper.status == "revision",
                (Paper.reviewed_by_id == current_admin.id) | (Paper.supervisor_id == current_admin.id)
            )
            .order_by(Paper.created_at.desc(), Paper.id.desc())
            .limit(200)
            .all()
        )
    elif reviewer_role == "project_coordinator":
        papers = (
            db.query(Paper)
            .join(User, Paper.created_by_id == User.id, isouter=True)
            .filter(Paper.status == "revision")
            .filter(func.lower(func.coalesce(User.department, "")) == reviewer_department)
            .order_by(Paper.created_at.desc(), Paper.id.desc())
            .limit(200)
            .all()
        )
    elif reviewer_role == "hod":
        papers = (
            db.query(Paper)
            .join(User, Paper.created_by_id == User.id, isouter=True)
            .filter(Paper.status == "revision")
            .filter(func.lower(func.coalesce(User.department, "")) == reviewer_department)
            .order_by(Paper.created_at.desc(), Paper.id.desc())
            .limit(200)
            .all()
        )
    else:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Reviewer role is not allowed")

    return [_to_paper_read(p) for p in papers]


@router.get("/papers/mine", response_model=list[PaperRead])
def read_my_papers(
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
) -> list[PaperRead]:
    papers = list_papers(
        db,
        status=None,
        sort="newest",
        limit=200,
        created_by_id=current_user.id,
    )
    return [_to_paper_read(p) for p in papers]


from fastapi import Request
from app.core.security import decode_access_token

@router.get("/papers/{paper_id}", response_model=PaperRead)
def read_paper(
    paper_id: int,
    request: Request,
    db: Session = Depends(get_db),
) -> PaperRead:
    paper = get_paper(db, paper_id)
    if not paper:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")
    
    current_user = None
    auth_header = request.headers.get("Authorization")
    if auth_header and auth_header.startswith("Bearer "):
        token = auth_header.split(" ")[1]
        try:
            payload = decode_access_token(token, expected_type="access")
            if payload and "sub" in payload:
                current_user = get_user_by_email(db, payload["sub"])
        except Exception:
            pass
            
    return _to_paper_read(paper, db, current_user)


@router.post("/papers", response_model=PaperRead, status_code=status.HTTP_201_CREATED)
def create_paper_endpoint(
    payload: PaperCreate,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
) -> PaperRead:
    if not _can_submit_paper(db, current_user):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Only students and lecturers can submit papers.",
        )
    if payload.department_id is None and current_user.department:
        matched_department_id = _safe_match_department_id(db, current_user.department)
        if matched_department_id is not None:
            payload = payload.model_copy(update={"department_id": matched_department_id})
    paper = create_paper(
        db,
        payload,
        created_by_id=current_user.id,
        created_by_is_admin=bool(current_user.is_admin),
    )
    lecturer_submission = _is_lecturer_submitter(db, current_user)
    if lecturer_submission and paper.status == "pending_lecturer":
        paper.status = "pending_hod"
        db.add(paper)
    _ensure_paper_doi(db, paper)
    _record_paper_version(
        db,
        paper=paper,
        source="student_upload" if not current_user.is_admin else "admin_upload",
        actor_id=current_user.id,
        note="Initial paper creation",
    )
    _record_workflow_event(
        db,
        paper_id=paper.id,
        event_type="paper_created",
        actor_id=current_user.id,
        actor_role=(current_user.role or "").strip().lower(),
        from_status=None,
        to_status=paper.status,
        message=f"Paper created: {paper.title}",
    )
    db.commit()
    db.refresh(paper)
    if not current_user.is_admin:
        if lecturer_submission:
            _notify_roles(
                db,
                roles={"hod"},
                paper_id=paper.id,
                department=(current_user.department or None),
                message=f"Lecturer research submitted and awaiting HOD approval: {paper.title}",
            )
            _notify_student(
                db,
                paper,
                f"Your paper '{paper.title}' was submitted and is awaiting HOD approval.",
            )
        else:
            _notify_assigned_lecturer(db, paper, current_user)
            _notify_hod_and_project_coordinator_on_student_submission(
                db,
                paper=paper,
                submitted_by=current_user,
            )
            _notify_student(
                db,
                paper,
                f"Your paper '{paper.title}' was submitted successfully and sent to lecturer review.",
            )
    return _to_paper_read(paper)


@router.post("/papers/upload", response_model=PaperRead, status_code=status.HTTP_201_CREATED)
async def upload_paper_endpoint(
    title: str = Form(...),
    abstract: str | None = Form(None),
    discipline: str | None = Form(None),
    university: str | None = Form(None),
    year: int | None = Form(None),
    document_type: str | None = Form(None),
    publication_type: str | None = Form("thesis"),
    license: str | None = Form(None),
    department_id: int | None = Form(None),
    supervisor_id: int | None = Form(None),
    work_mode: str | None = Form("individual"),
    tags: str | None = Form(None),
    authors: str | None = Form(None),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
) -> PaperRead:
    if not _can_submit_paper(db, current_user):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Only students and lecturers can submit papers.",
        )
    safe_name = f"{uuid4().hex}_{Path(file.filename or 'upload.bin').name}"
    dest = UPLOADS_DIR / safe_name

    content = await file.read()
    if not content:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Uploaded file is empty")
    dest.write_bytes(content)

    try:
        parsed_tags = json.loads(tags) if tags else []
        parsed_authors = json.loads(authors) if authors else []
        payload = PaperCreate(
            title=title,
            abstract=abstract,
            discipline=discipline,
            university=university,
            year=year,
            document_type=document_type,
            publication_type=(publication_type or "thesis").strip().lower(),
            license=license,
            tags=parsed_tags if isinstance(parsed_tags, list) else [],
            authors=parsed_authors if isinstance(parsed_authors, list) else [],
            file_name=file.filename,
            file_path=str(dest),
            file_size=len(content),
            mime_type=file.content_type,
            supervisor_id=supervisor_id,
            department_id=department_id,
            work_mode=(work_mode or "individual").strip().lower(),
        )
        if payload.department_id is None and current_user.department:
            matched_department_id = _safe_match_department_id(db, current_user.department)
            if matched_department_id is not None:
                payload = payload.model_copy(update={"department_id": matched_department_id})
    except Exception as exc:
        try:
            dest.unlink(missing_ok=True)
        except Exception:
            pass
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=f"Invalid upload metadata: {exc}") from exc

    paper = create_paper(
        db,
        payload,
        created_by_id=current_user.id,
        created_by_is_admin=bool(current_user.is_admin),
    )
    lecturer_submission = _is_lecturer_submitter(db, current_user)
    if lecturer_submission and paper.status == "pending_lecturer":
        paper.status = "pending_hod"
        db.add(paper)
    _ensure_paper_doi(db, paper)
    _record_paper_version(
        db,
        paper=paper,
        source="student_upload" if not current_user.is_admin else "admin_upload",
        actor_id=current_user.id,
        note="File upload",
    )
    _record_workflow_event(
        db,
        paper_id=paper.id,
        event_type="paper_uploaded",
        actor_id=current_user.id,
        actor_role=(current_user.role or "").strip().lower(),
        from_status=None,
        to_status=paper.status,
        message=f"File uploaded: {paper.file_name or paper.title}",
    )
    db.commit()
    db.refresh(paper)
    if not current_user.is_admin:
        if lecturer_submission:
            _notify_roles(
                db,
                roles={"hod"},
                paper_id=paper.id,
                department=(current_user.department or None),
                message=f"Lecturer research submitted and awaiting HOD approval: {paper.title}",
            )
            _notify_student(
                db,
                paper,
                f"Your paper '{paper.title}' was submitted and is awaiting HOD approval.",
            )
        else:
            _notify_assigned_lecturer(db, paper, current_user)
            _notify_hod_and_project_coordinator_on_student_submission(
                db,
                paper=paper,
                submitted_by=current_user,
            )
            _notify_student(
                db,
                paper,
                f"Your paper '{paper.title}' was submitted successfully and sent to lecturer review.",
            )
    return _to_paper_read(paper)


@router.post("/papers/{paper_id}/review", response_model=PaperRead)
def review_paper_endpoint(
    paper_id: int,
    payload: PaperReview,
    db: Session = Depends(get_db),
    current_admin=Depends(get_current_reviewer),
) -> PaperRead:
    paper = get_paper(db, paper_id)
    if not paper:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")
    reviewer_role = _resolve_reviewer_role(db, current_admin)

    if payload.decision != "approve":
        if reviewer_role in {"lecturer", "project_supervisor"} and payload.decision == "revision":
            latest_corrected = (
                db.query(PaperVersion)
                .filter(
                    PaperVersion.paper_id == paper.id,
                    PaperVersion.source == "supervisor_corrected_upload",
                )
                .order_by(PaperVersion.version_no.desc(), PaperVersion.id.desc())
                .first()
            )
            if not latest_corrected:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="Upload corrected file first before requesting revision.",
                )
        from_status = paper.status
        next_status = "revision" if payload.decision == "revision" else "rejected"
        updated = review_paper(db, paper, payload, reviewer_id=current_admin.id, next_status=next_status)
        _record_review_log(
            db,
            paper_id=updated.id,
            reviewer_id=current_admin.id,
            reviewer_role=reviewer_role,
            decision=payload.decision,
            comments=payload.comments,
            from_status=from_status,
            to_status=updated.status,
        )
        _record_workflow_event(
            db,
            paper_id=updated.id,
            event_type="review_decision",
            actor_id=current_admin.id,
            actor_role=reviewer_role,
            from_status=from_status,
            to_status=updated.status,
            message=payload.comments,
        )
        db.commit()
        db.refresh(updated)
        _notify_student(
            db,
            updated,
            f"Your paper '{updated.title}' was {next_status}. Feedback: {payload.comments or 'No feedback provided.'}",
        )
        return _to_paper_read(updated)

    if reviewer_role in {"lecturer", "project_supervisor"}:
        if paper.status != "pending_lecturer":
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Paper is not awaiting lecturer review")
        if paper.supervisor_id and paper.supervisor_id != current_admin.id:
            raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Only assigned lecturer can review this paper")
        if not paper.supervisor_id and paper.created_by_id == current_admin.id:
            raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="You cannot review your own unassigned paper")

        from_status = paper.status
        updated = review_paper(
            db,
            paper,
            payload,
            reviewer_id=current_admin.id,
            next_status="pending_hod_and_coordinator",
        )
        now = datetime.now(timezone.utc)
        updated.lecturer_approved_by_id = current_admin.id
        updated.lecturer_approved_at = now
        db.add(updated)
        _record_review_log(
            db,
            paper_id=updated.id,
            reviewer_id=current_admin.id,
            reviewer_role=reviewer_role,
            decision=payload.decision,
            comments=payload.comments,
            from_status=from_status,
            to_status=updated.status,
        )
        _record_workflow_event(
            db,
            paper_id=updated.id,
            event_type="workflow_transition",
            actor_id=current_admin.id,
            actor_role=reviewer_role,
            from_status=from_status,
            to_status=updated.status,
            message="Lecturer approval completed",
        )
        db.commit()
        db.refresh(updated)

        _notify_roles(
            db,
            roles={"hod", "project_coordinator"},
            paper_id=updated.id,
            department=(updated.created_by.department if updated.created_by else None),
            message=(
                "Dear Reviewer,\n\n"
                "This is to formally notify you that the following student submission has completed lecturer review "
                "and now requires your approval.\n\n"
                f"Title: {updated.title}\n"
                f"Student: {(updated.created_by.full_name if updated.created_by and updated.created_by.full_name else (updated.created_by.email if updated.created_by else 'Unknown Student'))}\n"
                f"Lecturer: {(current_admin.full_name or current_admin.email)}\n\n"
                "Please log in to MURRS to review and take the necessary action.\n\n"
                "Regards,\n"
                "MURRS Workflow System"
            ),
        )
        _notify_student(
            db,
            updated,
            f"Your paper '{updated.title}' passed lecturer review and is now awaiting HOD and Project Coordinator approvals.",
        )
        return _to_paper_read(updated)

    if reviewer_role == "project_coordinator":
        if paper.status not in {"pending_hod_and_coordinator", "pending_coordinator", "phase5_pending_coordinator"}:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Paper is not awaiting project coordinator review")
        if paper.project_coordinator_approved_at and paper.status != "phase5_pending_coordinator":
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Project coordinator approval already completed")

        from_status = paper.status
        updated = review_paper(
            db,
            paper,
            payload,
            reviewer_id=current_admin.id,
            next_status="pending_hod_and_coordinator" if paper.status != "phase5_pending_coordinator" else "phase5_pending_coordinator",
        )
        updated.project_coordinator_approved_by_id = current_admin.id
        updated.project_coordinator_approved_at = datetime.now(timezone.utc)
        if updated.status == "phase5_pending_coordinator":
            updated.status = "phase5_pending_hod"
        elif updated.hod_approved_at:
            updated.status = "approved_for_library"
        db.add(updated)
        _record_review_log(
            db,
            paper_id=updated.id,
            reviewer_id=current_admin.id,
            reviewer_role=reviewer_role,
            decision=payload.decision,
            comments=payload.comments,
            from_status=from_status,
            to_status=updated.status,
        )
        _record_workflow_event(
            db,
            paper_id=updated.id,
            event_type="workflow_transition",
            actor_id=current_admin.id,
            actor_role=reviewer_role,
            from_status=from_status,
            to_status=updated.status,
            message="Project coordinator approval completed",
        )
        db.commit()
        db.refresh(updated)

        if updated.status == "approved_for_library":
            _notify_supervisor_proposal_summary(db, updated)
            _notify_roles(
                db,
                roles={"librarian"},
                paper_id=updated.id,
                message=f"Paper fully approved by HOD and Project Coordinator. Ready for publication: {updated.title}",
            )
            _notify_student(
                db,
                updated,
                f"Your paper '{updated.title}' was approved by HOD and Project Coordinator, and has been sent to the librarian for publishing.",
            )
        elif from_status == "phase5_pending_coordinator":
            _notify_roles(
                db,
                roles={"hod"},
                paper_id=updated.id,
                department=(updated.created_by.department if updated.created_by else None),
                message=f"Paper corrections awaiting HOD final sign-off: {updated.title}",
            )
            _notify_student(
                db,
                updated,
                f"Your paper '{updated.title}' final corrections were signed off by the Project Coordinator and are awaiting final HOD sign-off.",
            )
        else:
            _notify_supervisor_proposal_summary(db, updated)
            _notify_roles(
                db,
                roles={"hod"},
                paper_id=updated.id,
                department=(updated.created_by.department if updated.created_by else None),
                message=f"Paper awaiting your HOD approval: {updated.title}",
            )
            _notify_student(
                db,
                updated,
                f"Your paper '{updated.title}' was approved by the Project Coordinator and is awaiting HOD approval.",
            )
        return _to_paper_read(updated)

    if reviewer_role == "hod":
        if paper.status not in {"pending_hod_and_coordinator", "pending_hod", "phase5_pending_hod"}:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Paper is not awaiting HOD review")
        if paper.hod_approved_at and paper.status != "phase5_pending_hod":
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="HOD approval already completed")

        from_status = paper.status
        updated = review_paper(
            db,
            paper,
            payload,
            reviewer_id=current_admin.id,
            next_status=paper.status,
        )
        updated.hod_approved_by_id = current_admin.id
        updated.hod_approved_at = datetime.now(timezone.utc)
        if updated.status == "phase5_pending_hod":
            updated.status = "approved_for_library"
        elif updated.project_coordinator_approved_at:
            updated.status = "approved_for_library"
        elif updated.created_by_id:
            owner = db.query(User).filter(User.id == updated.created_by_id).first()
            if owner and has_role(db, owner, "lecturer"):
                updated.status = "approved_for_library"
        db.add(updated)
        _record_review_log(
            db,
            paper_id=updated.id,
            reviewer_id=current_admin.id,
            reviewer_role=reviewer_role,
            decision=payload.decision,
            comments=payload.comments,
            from_status=from_status,
            to_status=updated.status,
        )
        _record_workflow_event(
            db,
            paper_id=updated.id,
            event_type="workflow_transition",
            actor_id=current_admin.id,
            actor_role=reviewer_role,
            from_status=from_status,
            to_status=updated.status,
            message="HOD approval completed",
        )
        db.commit()
        db.refresh(updated)

        if updated.status == "approved_for_library":
            _notify_supervisor_proposal_summary(db, updated)
            _notify_roles(
                db,
                roles={"librarian"},
                paper_id=updated.id,
                message=f"Paper fully approved by HOD and Project Coordinator. Ready for publication: {updated.title}",
            )
            _notify_student(
                db,
                updated,
                f"Your paper '{updated.title}' was approved by HOD and Project Coordinator, and has been sent to the librarian for publishing.",
            )
        else:
            _notify_supervisor_proposal_summary(db, updated)
            _notify_roles(
                db,
                roles={"project_coordinator"},
                paper_id=updated.id,
                department=(updated.created_by.department if updated.created_by else None),
                message=f"Paper awaiting your Project Coordinator approval: {updated.title}",
            )
            _notify_student(
                db,
                updated,
                f"Your paper '{updated.title}' was approved by the HOD and is awaiting Project Coordinator approval.",
            )
        return _to_paper_read(updated)

    if reviewer_role == "librarian":
        if paper.status != "approved_for_library":
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Paper is not awaiting librarian publishing")
        if payload.decision != "approve":
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Librarian can only publish work.",
            )
        from_status = paper.status

        # Librarian publishing must publish a PDF version.
        current_path = Path(paper.file_path or "")
        pdf_path = _convert_document_to_pdf_or_fail(current_path)
        paper.file_path = str(pdf_path)
        paper.file_name = f"{Path(paper.file_name or current_path.name).stem}.pdf"
        paper.mime_type = "application/pdf"
        paper.file_size = pdf_path.stat().st_size
        _record_paper_version(
            db,
            paper=paper,
            source="librarian_pdf_publish",
            actor_id=current_admin.id,
            note="Converted to PDF during librarian publish step",
        )
        db.add(paper)

        updated = review_paper(db, paper, payload, reviewer_id=current_admin.id, next_status="approved")
        _ensure_paper_doi(db, updated)
        _record_review_log(
            db,
            paper_id=updated.id,
            reviewer_id=current_admin.id,
            reviewer_role=reviewer_role,
            decision=payload.decision,
            comments=payload.comments,
            from_status=from_status,
            to_status=updated.status,
        )
        _record_workflow_event(
            db,
            paper_id=updated.id,
            event_type="workflow_transition",
            actor_id=current_admin.id,
            actor_role=reviewer_role,
            from_status=from_status,
            to_status=updated.status,
            message="Librarian publishing approval completed",
        )
        db.commit()
        db.refresh(updated)
        _notify_student(
            db,
            updated,
            f"Your paper '{updated.title}' has been published by the librarian and is now available online.",
        )
        return _to_paper_read(updated)

    raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Reviewer role is not allowed")


@router.post("/papers/{paper_id}/view", response_model=PaperRead)
def track_paper_view(paper_id: int, db: Session = Depends(get_db)) -> PaperRead:
    paper = get_paper(db, paper_id)
    if not paper:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")
    return _to_paper_read(increment_view(db, paper))


@router.post("/papers/{paper_id}/download", response_model=PaperRead)
def track_paper_download(
    paper_id: int,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
) -> PaperRead:
    paper = get_paper(db, paper_id)
    if not paper:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")
    return _to_paper_read(increment_download(db, paper))


@router.api_route("/papers/{paper_id}/binary", methods=["GET", "HEAD"])
def download_paper_file(
    paper_id: int,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    paper = get_paper(db, paper_id)
    if not paper:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")
    if not paper.file_path:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="No file attached to this paper")

    # ---------------------------------------------------------------
    # Phase-aware file selection:
    # Phase 2+ means the student has submitted a Proposal — serve the
    # latest accepted proposal file (or latest submitted if not yet accepted).
    # Phase 1 (topic submission) always serves the original file_path.
    # ---------------------------------------------------------------
    PROPOSAL_PHASES = {
        "phase2_proposal_submitted", "phase2_proposal_accepted",
        "phase3_chapters", "phase3_steps_in_progress", "phase3_all_steps_approved",
        "phase4_pending_examiners", "phase4_marking", "phase4_examination_completed",
        "phase5_corrections", "phase5_pending_supervisor", "phase5_pending_coordinator",
        "phase5_pending_hod", "phase5_pending_hod_and_coordinator",
        "phase5_approved_for_library", "phase5_published",
    }

    resolved_file_path = paper.file_path
    resolved_file_name = paper.file_name

    if paper.status in PROPOSAL_PHASES:
        from app.models.thesis_system import Proposal
        # Prefer accepted proposal, fallback to latest submitted
        accepted_proposal = (
            db.query(Proposal)
            .filter(Proposal.thesis_id == paper.id, Proposal.status == "accepted")
            .order_by(Proposal.version.desc())
            .first()
        )
        latest_proposal = accepted_proposal or (
            db.query(Proposal)
            .filter(Proposal.thesis_id == paper.id)
            .order_by(Proposal.version.desc())
            .first()
        )
        if latest_proposal and latest_proposal.file_url:
            proposal_path = Path(latest_proposal.file_url)
            if proposal_path.exists() and proposal_path.is_file():
                resolved_file_path = str(proposal_path)
                resolved_file_name = f"Proposal_v{latest_proposal.version}_{paper.title}.docx"

    path = Path(resolved_file_path)
    if not path.exists() or not path.is_file():
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Stored file not found")

    # Compile any database annotations/comments into the downloaded docx document
    from app.services.annotation_service import get_paper_annotations, compile_comments_to_docx
    annotations = get_paper_annotations(db, paper_id)
    download_path = Path(compile_comments_to_docx(resolved_file_path, annotations))

    download_name = resolved_file_name or download_path.name

    increment_download(db, paper)
    return FileResponse(
        path=download_path,
        media_type=paper.mime_type or "application/octet-stream",
        filename=download_name,
    )


@router.api_route("/papers/{paper_id}/proposal-file", methods=["GET", "HEAD"])
def download_proposal_file(
    paper_id: int,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    """Download the latest proposal file for a thesis (Phase 2+)."""
    paper = get_paper(db, paper_id)
    if not paper:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")

    from app.models.thesis_system import Proposal
    latest_proposal = (
        db.query(Proposal)
        .filter(Proposal.thesis_id == paper.id)
        .order_by(Proposal.version.desc())
        .first()
    )
    if not latest_proposal or not latest_proposal.file_url:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="No proposal file found for this thesis")

    path = Path(latest_proposal.file_url)
    if not path.exists() or not path.is_file():
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Proposal file not found on disk")

    download_name = f"Proposal_v{latest_proposal.version}_{paper.title}.docx"
    return FileResponse(
        path=path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=download_name,
    )


@router.api_route("/papers/{paper_id}/file", methods=["GET", "HEAD"])
def download_paper_file_legacy(
    paper_id: int,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    paper = get_paper(db, paper_id)
    if not paper:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")
    if not paper.file_path:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="No file attached to this paper")

    path = Path(paper.file_path)
    if not path.exists() or not path.is_file():
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Stored file not found")

    # Compile any database annotations/comments into the downloaded docx document
    from app.services.annotation_service import get_paper_annotations, compile_comments_to_docx
    annotations = get_paper_annotations(db, paper_id)
    download_path = Path(compile_comments_to_docx(paper.file_path, annotations))

    # Always return the currently stored file name so students receive
    # the exact supervisor-uploaded corrected file name.
    download_name = paper.file_name or download_path.name

    increment_download(db, paper)
    return FileResponse(
        path=download_path,
        media_type=paper.mime_type or "application/octet-stream",
        filename=download_name,
    )


@router.api_route("/papers/{paper_id}/reviewed-file", methods=["GET", "HEAD"])
def download_latest_reviewed_file(
    paper_id: int,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    paper = get_paper(db, paper_id)
    if not paper:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")

    allowed = bool(
        current_user.is_admin
        or current_user.id == paper.created_by_id
        or current_user.id == paper.supervisor_id
        or current_user.id == paper.reviewed_by_id
    )
    if not allowed:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Not allowed to access this file")

    latest_reviewed = (
        db.query(PaperVersion)
        .filter(
            PaperVersion.paper_id == paper_id,
            PaperVersion.source == "supervisor_corrected_upload",
        )
        .order_by(PaperVersion.version_no.desc(), PaperVersion.id.desc())
        .first()
    )

    from app.services.annotation_service import get_paper_annotations, compile_comments_to_docx
    annotations = get_paper_annotations(db, paper_id)

    source_file_path = None
    media_type = "application/octet-stream"
    download_name = None

    if latest_reviewed:
        source_file_path = latest_reviewed.file_path
        media_type = latest_reviewed.mime_type or media_type
        download_name = latest_reviewed.file_name
    elif annotations:
        source_file_path = paper.file_path
        media_type = paper.mime_type or media_type
        download_name = paper.file_name

    if not source_file_path:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="No supervisor reviewed file or comments available yet")

    path = Path(source_file_path)
    if not path.exists() or not path.is_file():
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Reviewed file not found on server")

    # Compile database annotations/comments into the downloaded docx document
    download_path = Path(compile_comments_to_docx(str(path), annotations))
    
    if not download_name:
        download_name = download_path.name

    increment_download(db, paper)
    return FileResponse(
        path=download_path,
        media_type=media_type,
        filename=download_name,
    )


@router.get("/papers/{paper_id}/feedback-file")
def download_paper_feedback_file(
    paper_id: int,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    paper = get_paper(db, paper_id)
    if not paper:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")

    # Author, assigned supervisor, reviewers, and admins can access feedback text.
    allowed = bool(
        current_user.is_admin
        or current_user.id == paper.created_by_id
        or current_user.id == paper.supervisor_id
        or current_user.id == paper.reviewed_by_id
    )
    if not allowed:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Not allowed to access feedback for this paper")

    from app.services.annotation_service import get_paper_annotations

    annotations = get_paper_annotations(db, paper_id)
    lines = [
        f"Paper: {paper.title or f'paper-{paper.id}'}",
        f"Status: {paper.status or '-'}",
        "",
        "Supervisor Review Comment:",
        paper.review_comments or "No review comment provided.",
        "",
        "Supervisor Annotation Notes:",
    ]
    if not annotations:
        lines.append("No annotation notes found.")
    else:
        for idx, a in enumerate(annotations, start=1):
            when = a.created_at.isoformat() if a.created_at else "-"
            lines.extend(
                [
                    f"{idx}. Location: {a.location or 'General note'}",
                    f"   Note: {a.text}",
                    f"   Created: {when}",
                    "",
                ]
            )

    content = "\n".join(lines)
    base_name = (paper.title or f"paper-{paper.id}").strip()
    file_name = f"{base_name}-feedback.txt"
    return Response(
        content=content,
        media_type="text/plain; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="{file_name}"'},
    )


def _build_multi_cert_excel_workbook(target_path: Path, sheet_prefix: str, papers_list: list) -> None:
    """Builds an ONLYOFFICE Excel workbook where each Certification Type has its own dedicated worksheet tab."""
    import openpyxl
    wb = openpyxl.Workbook()
    
    def classify_cert(p):
        dt = (str(getattr(p, "document_type", "") or "") + " " + str(getattr(p, "publication_type", "") or "") + " " + str(getattr(p, "discipline", "") or "")).lower()
        if "phd" in dt or "doctor" in dt:
            return "PhD"
        elif "mphil" in dt:
            return "MPhil"
        elif "master" in dt or "msc" in dt or "mba" in dt or "ma " in dt or "med" in dt:
            return "Masters (MSc-MBA)"
        else:
            return "Undergraduate (BSc-BA)"

    headers = [
        "Paper ID", "Certification Type", "Student Name", "Thesis Title", "Discipline / Dept",
        "Current Status", "Internal Score (0-100)", "External Score (0-100)",
        "Overall Grade / Rec", "General Comments"
    ]

    categories = [
        (f"All {sheet_prefix} Results", papers_list),
        ("Undergraduate (BSc-BA)", [p for p in papers_list if classify_cert(p) == "Undergraduate (BSc-BA)"]),
        ("Masters (MSc-MBA)", [p for p in papers_list if classify_cert(p) == "Masters (MSc-MBA)"]),
        ("MPhil", [p for p in papers_list if classify_cert(p) == "MPhil"]),
        ("PhD", [p for p in papers_list if classify_cert(p) == "PhD"]),
    ]

    first_sheet = True
    for cat_name, cat_papers in categories:
        if first_sheet:
            ws = wb.active
            ws.title = cat_name[:31]
            first_sheet = False
        else:
            ws = wb.create_sheet(title=cat_name[:31])

        ws.append(headers)
        for p in cat_papers:
            s_name = p.authors[0].name if p.authors else "Student"
            c_type = classify_cert(p)
            dept_name = p.discipline or (p.department.name if getattr(p, "department", None) else "")
            ws.append([
                p.id,
                c_type,
                s_name,
                p.title,
                dept_name,
                p.status,
                p.internal_score if p.internal_score is not None else "",
                p.external_score if p.external_score is not None else "",
                "Pass",
                p.review_comments or ""
            ])
            
    wb.save(target_path)


def _get_target_doc_path(paper, doc_type: str, user=None, assigned_papers=None, db: Session | None = None) -> tuple[Path, str, str, str]:
    """Returns (file_path, file_name, file_ext, document_type) for ONLYOFFICE."""
    base_dir = Path(paper.file_path).parent if paper and paper.file_path else Path("uploads/examiners")
    base_dir.mkdir(parents=True, exist_ok=True)
    
    student_name = "Student"
    if paper and paper.authors:
        student_name = paper.authors[0].name

    if doc_type == "comments" and paper:
        target = base_dir / f"paper_{paper.id}_comments.docx"
        file_display_name = f"Paper_{paper.id}_Examiner_Comments.docx"

        rebuild_needed = not target.exists()
        if target.exists():
            try:
                import docx
                existing_paragraphs = [p.text for p in docx.Document(target).paragraphs if p.text.strip()]
                existing_fulltext = "\n".join(existing_paragraphs)
                if "Examiner Evaluation Summary:" in existing_fulltext or "No detailed qualitative text notes entered yet" in existing_fulltext or (paper and paper.examiner_corrections and paper.examiner_corrections.strip() not in existing_fulltext):
                    rebuild_needed = True
            except Exception:
                rebuild_needed = True

        if rebuild_needed:
            try:
                import docx
                doc = docx.Document()
                doc.add_heading("Examiner Review & Qualitative Comments", level=1)
                doc.add_paragraph(f"Paper ID: #{paper.id}")
                doc.add_paragraph(f"Thesis Title: {paper.title}")
                doc.add_paragraph(f"Student Author: {student_name}")
                doc.add_paragraph("--------------------------------------------------------------------------------")
                doc.add_heading("Qualitative Feedback & Required Revisions:", level=2)
                
                int_comments = ""
                ext_comments = ""
                if db:
                    from app.models.thesis_system import ExaminationResult
                    results = db.query(ExaminationResult).filter(ExaminationResult.thesis_id == paper.id).all()
                    for r in results:
                        if r.examiner_type == "internal" or r.examiner_id == paper.internal_examiner_id:
                            if r.general_comments and r.general_comments.strip():
                                int_comments = r.general_comments.strip()
                        elif r.examiner_type == "external" or r.examiner_id == paper.external_examiner_id:
                            if r.general_comments and r.general_comments.strip():
                                ext_comments = r.general_comments.strip()

                has_notes = False
                if int_comments:
                    doc.add_paragraph(f"Internal Examiner Remarks:\n{int_comments}")
                    has_notes = True
                if ext_comments and ext_comments != int_comments:
                    doc.add_paragraph(f"External Examiner Remarks:\n{ext_comments}")
                    has_notes = True
                
                raw_corrections = _clean_examiner_corrections(paper.examiner_corrections)
                if raw_corrections and raw_corrections.strip() and raw_corrections.strip() != int_comments and raw_corrections.strip() != ext_comments:
                    doc.add_paragraph(f"Compiled Remarks & Instructions:\n{raw_corrections.strip()}")
                    has_notes = True

                if not has_notes:
                    doc.add_paragraph("")
                    doc.add_paragraph("")
                
                doc.save(target)
            except Exception:
                target.write_bytes(b"Examiner Comments Document\n")
        return target, file_display_name, "docx", "word"

    elif doc_type == "excel" and paper:
        is_internal = user and paper.internal_examiner_id == user.id
        is_external = user and paper.external_examiner_id == user.id
        role_label = "Internal Examiner" if is_internal else ("External Examiner" if is_external else "Examiner")
        role_tag = "internal" if is_internal else ("external" if is_external else f"user_{user.id}" if user else "gen")
        
        target = base_dir / f"paper_{paper.id}_marks_sheet_{role_tag}.xlsx"
        file_display_name = f"Paper_{paper.id}_{role_label.replace(' ', '_')}_Marks_Sheet.xlsx"

        if not target.exists():
            try:
                import openpyxl
                wb = openpyxl.Workbook()
                ws = wb.active
                ws.title = "Paper Marks Sheet"
                ws.append(["Paper ID", "Student Name", "Thesis Title", "Internal Score (0-100)", "External Score (0-100)", "Overall Recommendation", "Remarks"])
                ws.append([
                    paper.id,
                    student_name,
                    paper.title,
                    paper.internal_score if paper.internal_score is not None else "",
                    paper.external_score if paper.external_score is not None else "",
                    "Pass",
                    "Evaluation sheet ready"
                ])
                wb.save(target)
            except Exception:
                target.write_bytes(b"Paper Marks Sheet\n")
        return target, file_display_name, "xlsx", "cell"

    elif doc_type == "batch_excel" and user:
        upload_dir = Path("uploads/examiners")
        upload_dir.mkdir(parents=True, exist_ok=True)
        target = upload_dir / f"assigned_results_user_{user.id}.xlsx"
        try:
            _build_multi_cert_excel_workbook(target, "Assigned", assigned_papers or [])
        except Exception:
            target.write_bytes(b"Assigned Results Sheet\n")
        return target, f"Assigned_Thesis_Results_{user.id}.xlsx", "xlsx", "cell"

    elif doc_type == "dept_excel" and user:
        upload_dir = Path("uploads/examiners")
        upload_dir.mkdir(parents=True, exist_ok=True)
        target = upload_dir / f"department_results_user_{user.id}.xlsx"
        try:
            _build_multi_cert_excel_workbook(target, "Department", assigned_papers or [])
        except Exception:
            target.write_bytes(b"Department Results Sheet\n")
        return target, f"Department_Thesis_Results_{user.id}.xlsx", "xlsx", "cell"

    elif doc_type == "dean_excel" and user:
        upload_dir = Path("uploads/examiners")
        upload_dir.mkdir(parents=True, exist_ok=True)
        target = upload_dir / f"dean_master_results_user_{user.id}.xlsx"
        try:
            _build_multi_cert_excel_workbook(target, "School Master", assigned_papers or [])
        except Exception:
            target.write_bytes(b"Dean School Master Results Sheet\n")
        return target, f"Dean_School_Master_Results_{user.id}.xlsx", "xlsx", "cell"

    # Default: student's supervisor-approved main thesis file
    target = Path(paper.file_path) if (paper and paper.file_path and Path(paper.file_path).exists()) else None
    if not target or not target.exists():
        target_dir = Path("uploads/theses")
        target_dir.mkdir(parents=True, exist_ok=True)
        target = target_dir / f"paper_{paper.id}_approved_thesis.docx" if paper else Path("uploads/paper.docx")
        if not target.exists() and paper:
            try:
                import docx
                doc = docx.Document()
                doc.add_heading(f"Approved Final Thesis: {paper.title}", level=1)
                doc.add_paragraph(f"Student Author: {student_name}")
                doc.add_paragraph(f"Paper ID: #{paper.id}")
                doc.add_paragraph(f"Status: {paper.status} (Supervisor Signed Off)")
                doc.add_paragraph("--------------------------------------------------------------------------------")
                doc.add_heading("Abstract & Full Thesis Submission", level=2)
                doc.add_paragraph(paper.abstract or "Full thesis work submitted by student and approved by supervisor for Phase 4 examination.")
                doc.save(target)
            except Exception:
                target.write_bytes(b"Student Approved Thesis Document\n")

    file_name = (paper.file_name if (paper and paper.file_name) else None) or target.name
    file_ext = (target.suffix or "").lstrip(".").lower() or "docx"
    doc_kind = "cell" if file_ext in {"xlsx", "xls", "csv"} else ("slide" if file_ext in {"pptx", "ppt"} else "word")
    return target, file_name, file_ext, doc_kind


@router.api_route("/papers/{paper_id}/file/public", methods=["GET", "HEAD"])
def download_paper_file_public(
    paper_id: int,
    token: str = Query(..., min_length=20),
    doc_type: str = Query("paper"),
    uid: int | None = Query(None),
    db: Session = Depends(get_db),
):
    if not _verify_editor_token(token=token, paper_id=paper_id, action="file"):
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Invalid or expired file token")

    paper = get_paper(db, paper_id) if paper_id > 0 else None
    if not paper and doc_type not in {"batch_excel", "dept_excel", "dean_excel"}:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")

    user = db.query(User).filter(User.id == uid).first() if uid else None
    assigned_papers = []
    if user:
        if doc_type == "batch_excel":
            assigned_papers = db.query(Paper).filter(
                (Paper.internal_examiner_id == user.id) | (Paper.external_examiner_id == user.id) | (Paper.supervisor_id == user.id)
            ).all()
            if not assigned_papers:
                assigned_papers = db.query(Paper).all()
        elif doc_type == "dept_excel":
            u_dept = (user.department or "").strip().lower()
            if u_dept:
                assigned_papers = (
                    db.query(Paper)
                    .join(User, Paper.created_by_id == User.id, isouter=True)
                    .filter(
                        (func.lower(func.coalesce(Paper.discipline, "")) == u_dept)
                        | (func.lower(func.coalesce(User.department, "")) == u_dept)
                    )
                    .all()
                )
            else:
                assigned_papers = db.query(Paper).all()
        elif doc_type == "dean_excel":
            assigned_papers = db.query(Paper).all()

    target, file_name, file_ext, _ = _get_target_doc_path(paper, doc_type, user=user, assigned_papers=assigned_papers, db=db)
    if not target.exists() or not target.is_file():
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Stored file not found")

    mime_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet" if file_ext == "xlsx" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return FileResponse(
        path=target,
        media_type=mime_type,
        filename=file_name,
    )


@router.get("/papers/examiner/results-excel-config")
def get_examiner_results_excel_config(
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    if not settings.onlyoffice_doc_server_url:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="OnlyOffice is not configured")

    assigned_query = db.query(Paper).filter(
        (Paper.internal_examiner_id == current_user.id) | (Paper.external_examiner_id == current_user.id) | (Paper.supervisor_id == current_user.id)
    ).all()
    if not assigned_query:
        assigned_query = db.query(Paper).all()

    target, file_name, file_ext, doc_kind = _get_target_doc_path(None, "batch_excel", user=current_user, assigned_papers=assigned_query)

    file_token = urllib.parse.quote(_build_editor_token(paper_id=0, action="file"), safe="")
    callback_token = urllib.parse.quote(_build_editor_token(paper_id=0, action="callback"), safe="")
    callback_base = (settings.onlyoffice_callback_base_url or settings.public_api_base_url).rstrip("/")
    file_url = f"{callback_base}{settings.api_prefix}/papers/0/file/public?token={file_token}&doc_type=batch_excel&uid={current_user.id}"
    callback_url = f"{callback_base}{settings.api_prefix}/papers/0/editor-callback?token={callback_token}&doc_type=batch_excel&uid={current_user.id}"

    session_key = f"examiner-excel-{current_user.id}-{target.stat().st_size if target.exists() else 0}-{int(datetime.now(timezone.utc).timestamp())}"

    config = {
        "documentType": "cell",
        "type": "desktop",
        "document": {
            "title": file_name,
            "url": file_url,
            "fileType": "xlsx",
            "key": session_key,
        },
        "editorConfig": {
            "callbackUrl": callback_url,
            "mode": "edit",
            "lang": "en",
            "user": {
                "id": str(current_user.id),
                "name": current_user.full_name or current_user.email,
            },
        },
    }
    return {
        "document_server_url": settings.onlyoffice_doc_server_url.rstrip("/"),
        "config": config,
    }


@router.get("/papers/department/results-excel-config")
def get_department_results_excel_config(
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    if not settings.onlyoffice_doc_server_url:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="OnlyOffice is not configured")

    u_dept = (current_user.department or "").strip().lower()
    if u_dept:
        dept_papers = (
            db.query(Paper)
            .join(User, Paper.created_by_id == User.id, isouter=True)
            .filter(
                (func.lower(func.coalesce(Paper.discipline, "")) == u_dept)
                | (func.lower(func.coalesce(User.department, "")) == u_dept)
            )
            .all()
        )
    else:
        dept_papers = db.query(Paper).all()

    target, file_name, file_ext, doc_kind = _get_target_doc_path(None, "dept_excel", user=current_user, assigned_papers=dept_papers)

    file_token = urllib.parse.quote(_build_editor_token(paper_id=0, action="file"), safe="")
    callback_token = urllib.parse.quote(_build_editor_token(paper_id=0, action="callback"), safe="")
    callback_base = (settings.onlyoffice_callback_base_url or settings.public_api_base_url).rstrip("/")
    file_url = f"{callback_base}{settings.api_prefix}/papers/0/file/public?token={file_token}&doc_type=dept_excel&uid={current_user.id}"
    callback_url = f"{callback_base}{settings.api_prefix}/papers/0/editor-callback?token={callback_token}&doc_type=dept_excel&uid={current_user.id}"

    session_key = f"dept-excel-{current_user.id}-{target.stat().st_size if target.exists() else 0}-{int(datetime.now(timezone.utc).timestamp())}"

    config = {
        "documentType": "cell",
        "type": "desktop",
        "document": {
            "title": file_name,
            "url": file_url,
            "fileType": "xlsx",
            "key": session_key,
        },
        "editorConfig": {
            "callbackUrl": callback_url,
            "mode": "edit",
            "lang": "en",
            "user": {
                "id": str(current_user.id),
                "name": current_user.full_name or current_user.email,
            },
        },
    }
    return {
        "document_server_url": settings.onlyoffice_doc_server_url.rstrip("/"),
        "config": config,
    }


@router.get("/papers/dean/results-excel-config")
def get_dean_results_excel_config(
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    if not settings.onlyoffice_doc_server_url:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="OnlyOffice is not configured")

    school_papers = db.query(Paper).all()

    target, file_name, file_ext, doc_kind = _get_target_doc_path(None, "dean_excel", user=current_user, assigned_papers=school_papers)

    file_token = urllib.parse.quote(_build_editor_token(paper_id=0, action="file"), safe="")
    callback_token = urllib.parse.quote(_build_editor_token(paper_id=0, action="callback"), safe="")
    callback_base = (settings.onlyoffice_callback_base_url or settings.public_api_base_url).rstrip("/")
    file_url = f"{callback_base}{settings.api_prefix}/papers/0/file/public?token={file_token}&doc_type=dean_excel&uid={current_user.id}"
    callback_url = f"{callback_base}{settings.api_prefix}/papers/0/editor-callback?token={callback_token}&doc_type=dean_excel&uid={current_user.id}"

    session_key = f"dean-excel-{current_user.id}-{target.stat().st_size if target.exists() else 0}-{int(datetime.now(timezone.utc).timestamp())}"

    config = {
        "documentType": "cell",
        "type": "desktop",
        "document": {
            "title": file_name,
            "url": file_url,
            "fileType": "xlsx",
            "key": session_key,
        },
        "editorConfig": {
            "callbackUrl": callback_url,
            "mode": "edit",
            "lang": "en",
            "user": {
                "id": str(current_user.id),
                "name": current_user.full_name or current_user.email,
            },
        },
    }
    return {
        "document_server_url": settings.onlyoffice_doc_server_url.rstrip("/"),
        "config": config,
    }


@router.get("/papers/{paper_id}/editor-config")
def get_editor_config(
    paper_id: int,
    doc_type: str = Query("paper"),
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    if not settings.onlyoffice_doc_server_url:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="OnlyOffice is not configured")

    paper = get_paper(db, paper_id)
    if not paper:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")

    target, file_name, file_ext, doc_kind = _get_target_doc_path(paper, doc_type, user=current_user, db=db)

    # Basic access gate for editor config.
    if not current_user.is_admin and paper.created_by_id != current_user.id and paper.supervisor_id not in {None, current_user.id}:
        if not has_role(db, current_user, "librarian") and not has_role(db, current_user, "hod") and not has_role(db, current_user, "project_coordinator") and paper.internal_examiner_id != current_user.id and paper.external_examiner_id != current_user.id:
            raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Not allowed to access this document")

    file_token = urllib.parse.quote(_build_editor_token(paper_id=paper.id, action="file"), safe="")
    callback_token = urllib.parse.quote(_build_editor_token(paper_id=paper.id, action="callback"), safe="")
    callback_base = (settings.onlyoffice_callback_base_url or settings.public_api_base_url).rstrip("/")
    file_url = f"{callback_base}{settings.api_prefix}/papers/{paper.id}/file/public?token={file_token}&doc_type={doc_type}&uid={current_user.id}&v={int(target.stat().st_mtime if target.exists() else 0)}"
    callback_url = f"{callback_base}{settings.api_prefix}/papers/{paper.id}/editor-callback?token={callback_token}&doc_type={doc_type}&uid={current_user.id}"

    session_key = f"paper-{paper.id}-{doc_type}-{current_user.id}-{target.stat().st_size if target.exists() else 0}-{int(datetime.now(timezone.utc).timestamp())}"

    config = {
        "documentType": doc_kind,
        "type": "desktop",
        "document": {
            "title": file_name,
            "url": file_url,
            "fileType": file_ext,
            "key": session_key,
        },
        "editorConfig": {
            "callbackUrl": callback_url,
            "mode": "edit",
            "lang": "en",
            "user": {
                "id": str(current_user.id),
                "name": current_user.full_name or current_user.email,
            },
        },
    }
    return {
        "document_server_url": settings.onlyoffice_doc_server_url.rstrip("/"),
        "config": config,
    }


def _sync_batch_excel_to_db(target: Path, db: Session, user):
    try:
        import openpyxl
        wb = openpyxl.load_workbook(target, data_only=True)
        ws = wb.active
        for r in range(2, ws.max_row + 1):
            pid_val = ws.cell(row=r, column=1).value
            if pid_val is None:
                continue
            try:
                pid = int(str(pid_val).replace("PAPER-", "").replace("#", "").strip())
            except ValueError:
                continue
            
            p = get_paper(db, pid)
            if not p:
                continue
            
            int_score_cell = ws.cell(row=r, column=6).value
            ext_score_cell = ws.cell(row=r, column=7).value
            rec_cell = ws.cell(row=r, column=8).value
            comments_cell = ws.cell(row=r, column=9).value
            
            if int_score_cell is not None and str(int_score_cell).strip():
                try:
                    p.internal_score = float(int_score_cell)
                except ValueError:
                    pass
            if ext_score_cell is not None and str(ext_score_cell).strip():
                try:
                    p.external_score = float(ext_score_cell)
                except ValueError:
                    pass
            if comments_cell and str(comments_cell).strip():
                p.review_comments = str(comments_cell).strip()
            db.add(p)
        db.commit()
    except Exception:
        pass


def _sync_marks_sheet_to_db(target: Path, db: Session, paper):
    try:
        import openpyxl
        wb = openpyxl.load_workbook(target, data_only=True)
        ws = wb.active
        if ws.max_row >= 2:
            int_score_cell = ws.cell(row=2, column=4).value
            ext_score_cell = ws.cell(row=2, column=5).value
            rec_cell = ws.cell(row=2, column=6).value
            remarks_cell = ws.cell(row=2, column=7).value
            
            if int_score_cell is not None and str(int_score_cell).strip():
                try:
                    paper.internal_score = float(int_score_cell)
                except ValueError:
                    pass
            if ext_score_cell is not None and str(ext_score_cell).strip():
                try:
                    paper.external_score = float(ext_score_cell)
                except ValueError:
                    pass
            if remarks_cell and str(remarks_cell).strip():
                paper.review_comments = str(remarks_cell).strip()
            db.add(paper)
            db.commit()
    except Exception:
        pass


@router.post("/papers/{paper_id}/editor-callback")
async def handle_editor_callback(
    paper_id: int,
    token: str = Query(..., min_length=20),
    doc_type: str = Query("paper"),
    uid: int | None = Query(None),
    payload: dict | None = None,
    db: Session = Depends(get_db),
):
    if not _verify_editor_token(token=token, paper_id=paper_id, action="callback"):
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Invalid or expired callback token")

    paper = get_paper(db, paper_id) if paper_id > 0 else None
    user = db.query(User).filter(User.id == uid).first() if uid else None

    target, _, _, _ = _get_target_doc_path(paper, doc_type, user=user)

    data = payload or {}
    status_code = int(data.get("status", 0) or 0)
    download_url = data.get("url")
    if status_code in {2, 6} and isinstance(download_url, str) and download_url.strip():
        try:
            with urllib.request.urlopen(download_url, timeout=30) as response:
                content = response.read()
            if not content:
                return {"error": 1}
            target.write_bytes(content)
            if paper and doc_type == "paper":
                paper.file_size = len(content)
                db.add(paper)
                db.commit()
            elif doc_type == "batch_excel" and user:
                _sync_batch_excel_to_db(target, db, user)
            elif doc_type == "excel" and paper:
                _sync_marks_sheet_to_db(target, db, paper)
        except Exception:
            return {"error": 1}
    return {"error": 0}


@router.post("/papers/{paper_id}/corrected-file", response_model=PaperRead)
async def upload_corrected_file(
    paper_id: int,
    file: UploadFile = File(...),
    note: str | None = Form(None),
    db: Session = Depends(get_db),
    current_admin=Depends(get_current_reviewer),
) -> PaperRead:
    paper = get_paper(db, paper_id)
    if not paper:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")

    reviewer_role = _resolve_reviewer_role(db, current_admin)
    if reviewer_role not in {"lecturer", "project_supervisor"}:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Only lecturers or project supervisors can upload corrected files",
        )
    if paper.supervisor_id and paper.supervisor_id != current_admin.id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Only the assigned supervisor can upload corrected files",
        )

    content = await file.read()
    if not content:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Uploaded file is empty")

    previous_path = Path(paper.file_path) if paper.file_path else None
    note_text = (note or "").strip()

    if previous_path and previous_path.exists() and previous_path.is_file():
        old_hash = hashlib.sha256(previous_path.read_bytes()).hexdigest()
        new_hash = hashlib.sha256(content).hexdigest()
        if old_hash == new_hash:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Uploaded corrected file is identical to the current paper file. Please upload the edited file.",
            )

    safe_name = f"{uuid4().hex}_{Path(file.filename or 'corrected.bin').name}"
    dest = UPLOADS_DIR / safe_name
    dest.write_bytes(content)
    from_status = paper.status
    try:
        paper.file_name = file.filename or safe_name
        paper.file_path = str(dest)
        paper.file_size = len(content)
        paper.mime_type = file.content_type
        paper.status = "revision"
        if note_text:
            paper.review_comments = note_text

        db.add(paper)
        _record_paper_version(
            db,
            paper=paper,
            source="supervisor_corrected_upload",
            actor_id=current_admin.id,
            note=note_text or "Supervisor uploaded corrected file",
        )
        _record_workflow_event(
            db,
            paper_id=paper.id,
            event_type="corrected_file_uploaded",
            actor_id=current_admin.id,
            actor_role=reviewer_role,
            from_status=from_status,
            to_status=paper.status,
            message=note_text or "Corrected file uploaded by supervisor",
        )
        db.commit()
        db.refresh(paper)
    except Exception:
        try:
            dest.unlink(missing_ok=True)
        except Exception:
            pass
        raise

    if previous_path and previous_path != dest:
        try:
            previous_path.unlink(missing_ok=True)
        except Exception:
            pass

    _notify_student(
        db,
        paper,
        (
            f"Your supervisor uploaded a corrected version for '{paper.title}'. "
            f"{'Note: ' + note_text if note_text else 'Check the updated file and resubmit your final version.'}"
        ),
    )

    # Also email the student with the exact corrected file attached.
    subject = f"Corrected File Uploaded: {paper.title}"
    message = (
        f"Your supervisor uploaded a corrected file for '{paper.title}'.\n\n"
        f"{('Supervisor note: ' + note_text + '\\n\\n') if note_text else ''}"
        "The corrected file is attached to this email."
    )
    attachment_name = paper.file_name or Path(paper.file_path or "").name or "corrected-file.bin"
    attachment_mime = paper.mime_type or "application/octet-stream"
    recipients_sent: set[str] = set()

    if paper.created_by and paper.created_by.email:
        mail = paper.created_by.email.strip().lower()
        if mail:
            recipients_sent.add(mail)
            send_notification_email(
                to_email=mail,
                to_name=paper.created_by.full_name,
                subject=subject,
                message=message,
                attachments=[(attachment_name, content, attachment_mime)],
            )

    # Fallback for legacy/imported papers without created_by_id.
    for author in (paper.authors or []):
        author_email = (author.email or "").strip().lower() if author else ""
        if not author_email or author_email in recipients_sent:
            continue
        recipients_sent.add(author_email)
        send_notification_email(
            to_email=author_email,
            to_name=(author.name if author else None),
            subject=subject,
            message=message,
            attachments=[(attachment_name, content, attachment_mime)],
        )

    return _to_paper_read(paper)


@router.post("/papers/{paper_id}/annotations", status_code=status.HTTP_201_CREATED)
def create_annotation_endpoint(
    paper_id: int,
    location: str | None = Form(None),
    text: str = Form(...),
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    """Create an annotation on a paper. Supervisors can annotate; authors can view."""
    from app.services.annotation_service import create_annotation
    
    paper = get_paper(db, paper_id)
    if not paper:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")
    
    if not _is_annotation_participant(paper, current_user):
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Only supervisors and authors can annotate")
    
    annotation = create_annotation(db, paper_id, current_user.id, text, location)
    author_name = current_user.full_name or current_user.email or "Supervisor"
    initials = "".join([p[0].upper() for p in author_name.split() if p])[:2] or "DT"
    return {
        "id": annotation.id,
        "paper_id": annotation.paper_id,
        "author_id": annotation.author_id,
        "author_name": author_name,
        "author_initials": initials,
        "location": annotation.location,
        "text": annotation.text,
        "resolved": annotation.resolved,
        "created_at": annotation.created_at.isoformat() if annotation.created_at else None,
    }


@router.get("/papers/{paper_id}/annotations")
def get_annotations_endpoint(
    paper_id: int,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    """Get all annotations for a paper."""
    from app.services.annotation_service import get_paper_annotations
    
    paper = get_paper(db, paper_id)
    if not paper:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")
    
    if not _is_annotation_participant(paper, current_user):
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Only supervisors and authors can view annotations")
    
    annotations = get_paper_annotations(db, paper_id)
    res = []
    for a in annotations:
        a_name = (a.author.full_name or a.author.email) if a.author else "Supervisor"
        initials = "".join([p[0].upper() for p in a_name.split() if p])[:2] or "DT"
        res.append({
            "id": a.id,
            "paper_id": a.paper_id,
            "author_id": a.author_id,
            "author_name": a_name,
            "author_initials": initials,
            "location": a.location,
            "text": a.text,
            "resolved": a.resolved,
            "created_at": a.created_at.isoformat() if a.created_at else None,
        })
    return res


@router.patch("/papers/{paper_id}/annotations/{annotation_id}")
def update_annotation_endpoint(
    paper_id: int,
    annotation_id: int,
    text: str | None = Form(None),
    resolved: bool | None = Form(None),
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    """Update an annotation. Only the author can update."""
    from app.services.annotation_service import get_annotation, update_annotation
    
    paper = get_paper(db, paper_id)
    if not paper:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")
    
    annotation = get_annotation(db, annotation_id)
    if not annotation:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Annotation not found")
    
    if annotation.author_id != current_user.id and not current_user.is_admin:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Only the annotation author can update it")
    
    updated = update_annotation(db, annotation, text=text, resolved=resolved)
    return {
        "id": updated.id,
        "paper_id": updated.paper_id,
        "author_id": updated.author_id,
        "location": updated.location,
        "text": updated.text,
        "resolved": updated.resolved,
        "created_at": updated.created_at.isoformat() if updated.created_at else None,
    }


@router.delete("/papers/{paper_id}/annotations/{annotation_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_annotation_endpoint(
    paper_id: int,
    annotation_id: int,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    """Delete an annotation. Only the author can delete."""
    from app.services.annotation_service import get_annotation, delete_annotation
    
    paper = get_paper(db, paper_id)
    if not paper:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")
    
    annotation = get_annotation(db, annotation_id)
    if not annotation:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Annotation not found")
    
    if annotation.author_id != current_user.id and not current_user.is_admin:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Only the annotation author can delete it")
    
    delete_annotation(db, annotation_id)
    return None


@router.post("/papers/{paper_id}/supervisors")
def assign_supervisors_endpoint(
    paper_id: int,
    supervisor_user_ids: list[int] = Form(...),
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    """Assign supervisors to a paper. HOD or Admin only."""
    from app.services.annotation_service import assign_supervisors_to_paper
    from app.services.user_service import has_role
    
    paper = get_paper(db, paper_id)
    if not paper:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")
    
    # Check if user is HOD or Admin
    is_hod = has_role(db, current_user, "hod")
    if not (is_hod or current_user.is_admin):
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Only HOD or Admin can assign supervisors")
    if is_hod and not current_user.is_admin:
        paper_department = (paper.created_by.department if paper.created_by else None) or ""
        if (paper_department or "").strip().lower() != (current_user.department or "").strip().lower():
            raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="HOD can only assign supervisors for papers in their department")
    
    try:
        supervisors = assign_supervisors_to_paper(db, paper_id, supervisor_user_ids, assigned_by_id=current_user.id)
        return [
            {
                "id": s.id,
                "paper_id": s.paper_id,
                "user_id": s.user_id,
                "assigned_by_id": s.assigned_by_id,
                "created_at": s.created_at.isoformat() if s.created_at else None,
            }
            for s in supervisors
        ]
    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(e))


@router.post("/papers/{paper_id}/assign-coordinator", response_model=PaperRead)
def assign_coordinator(
    paper_id: int,
    project_coordinator_id: int = Form(...),
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
) -> PaperRead:
    paper = get_paper(db, paper_id)
    if not paper:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")
    
    # Only HOD or Admin can assign Project Coordinator
    if not (has_role(db, current_user, "hod") or current_user.is_admin):
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Only HOD or Admin can assign the Project Coordinator")
    
    coordinator_user = db.query(User).filter(User.id == project_coordinator_id).first()
    if not coordinator_user or not has_role(db, coordinator_user, "project_coordinator"):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Selected user must have the Project Coordinator role")
    
    paper.project_coordinator_id = project_coordinator_id
    paper.status = "phase2_pending_supervisor"
    db.add(paper)
    
    _record_workflow_event(
        db,
        paper_id=paper.id,
        event_type="assign_coordinator",
        actor_id=current_user.id,
        actor_role="hod",
        from_status="phase1_proposal_submitted",
        to_status=paper.status,
        message=f"Project Coordinator assigned: {coordinator_user.full_name or coordinator_user.email}",
    )
    db.commit()
    db.refresh(paper)
    
    _notify_hod_and_coordinators(
        db,
        paper,
        f"Project Coordinator Assigned: Coordinator {coordinator_user.full_name or coordinator_user.email} assigned to '{paper.title}'."
    )
    return _to_paper_read(paper, db, current_user)


@router.post("/papers/{paper_id}/assign-supervisor", response_model=PaperRead)
def assign_supervisor(
    paper_id: int,
    supervisor_id: int = Form(...),
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
) -> PaperRead:
    paper = get_paper(db, paper_id)
    if not paper:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")
    
    # Check permissions: Admin, or HOD/Coordinator of the same department
    is_admin = current_user.is_admin or has_role(db, current_user, "system_admin")
    is_hod = has_role(db, current_user, "hod")
    is_coordinator = has_role(db, current_user, "project_coordinator")
    
    if not is_admin:
        if not (is_hod or is_coordinator):
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Only the Project Coordinator, HOD, or Admin can assign the supervisor"
            )
        # Check department match
        user_dept = (current_user.department or "").strip().lower()
        paper_dept = ""
        if paper.department:
            paper_dept = (paper.department.name or "").strip().lower()
        elif paper.discipline:
            paper_dept = (paper.discipline or "").strip().lower()
            
        if not user_dept or not paper_dept or user_dept != paper_dept:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="You can only assign supervisors to papers within your own department"
            )

    supervisor_user = db.query(User).filter(User.id == supervisor_id).first()
    if not supervisor_user:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Selected supervisor not found")
    
    from app.models.paper import PaperSupervisor
    paper.supervisor_id = supervisor_id
    
    # Insert into PaperSupervisor relationship
    if not any(s.user_id == supervisor_id for s in paper.supervisors):
        db.add(PaperSupervisor(paper_id=paper.id, user_id=supervisor_id, assigned_by_id=current_user.id))
    
    from_status = paper.status
    paper.status = "phase1_topic_accepted"
    db.add(paper)
    
    _record_workflow_event(
        db,
        paper_id=paper.id,
        event_type="assign_supervisor",
        actor_id=current_user.id,
        actor_role="project_coordinator" if is_coordinator else "hod",
        from_status=from_status,
        to_status=paper.status,
        message=f"Supervisor assigned: {supervisor_user.full_name or supervisor_user.email}",
    )
    db.commit()
    db.refresh(paper)
    
    _notify_hod_and_coordinators(
        db,
        paper,
        f"Topic Approved & Supervisor Assigned: Supervisor {supervisor_user.full_name or supervisor_user.email} assigned to '{paper.title}'."
    )
    _notify_student(db, paper, f"Your thesis topic was approved! Supervisor {supervisor_user.full_name or supervisor_user.email} has been assigned to supervise your work. You can now submit your Project Proposal for Phase 2.")
    
    create_notification(
        db,
        user_id=supervisor_id,
        paper_id=paper.id,
        ntype="workflow_update",
        message=f"You have been assigned to supervise student paper: '{paper.title}'"
    )
    
    return _to_paper_read(paper, db, current_user)


@router.post("/papers/{paper_id}/chapters/student", response_model=PaperRead)
def student_update_checklist(
    paper_id: int,
    ch1: bool = Form(...),
    ch2: bool = Form(...),
    ch3: bool = Form(...),
    ch4: bool = Form(...),
    ch5: bool = Form(...),
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
) -> PaperRead:
    paper = get_paper(db, paper_id)
    if not paper:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")
    
    if paper.created_by_id != current_user.id and not current_user.is_admin:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Only the student author can update their checklist progress")
    
    paper.ch1_student_done = ch1
    paper.ch2_student_done = ch2
    paper.ch3_student_done = ch3
    paper.ch4_student_done = ch4
    paper.ch5_student_done = ch5
    db.add(paper)
    db.commit()
    db.refresh(paper)
    
    if paper.supervisor_id:
        create_notification(
            db,
            user_id=paper.supervisor_id,
            paper_id=paper.id,
            ntype="workflow_update",
            message=f"Student updated chapter checklist for '{paper.title}'. Please review and approve."
        )
        
    return _to_paper_read(paper, db, current_user)


@router.post("/papers/{paper_id}/chapters/supervisor", response_model=PaperRead)
def supervisor_update_checklist(
    paper_id: int,
    ch1: bool = Form(...),
    ch2: bool = Form(...),
    ch3: bool = Form(...),
    ch4: bool = Form(...),
    ch5: bool = Form(...),
    comments: str | None = Form(None),
    rejected_chapter: int | None = Form(None),
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
) -> PaperRead:
    paper = get_paper(db, paper_id)
    if not paper:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")
    
    if paper.supervisor_id != current_user.id and not current_user.is_admin:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Only the assigned supervisor can approve chapter progress")

    chapter_values = [ch1, ch2, ch3, ch4, ch5]
    for index, approved in enumerate(chapter_values[1:], start=2):
        if approved and not chapter_values[index - 2]:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Chapter {index} cannot be approved before Chapter {index - 1}",
            )
    
    paper.ch1_supervisor_approved = ch1
    paper.ch2_supervisor_approved = ch2
    paper.ch3_supervisor_approved = ch3
    paper.ch4_supervisor_approved = ch4
    paper.ch5_supervisor_approved = ch5
    if rejected_chapter:
        paper.status = "revision"
        # Reset the student done checkbox for the rejected chapter
        if rejected_chapter == 1:
            paper.ch1_student_done = False
        elif rejected_chapter == 2:
            paper.ch2_student_done = False
        elif rejected_chapter == 3:
            paper.ch3_student_done = False
        elif rejected_chapter == 4:
            paper.ch4_student_done = False
        elif rejected_chapter == 5:
            paper.ch5_student_done = False
    else:
        # If the supervisor just approved a chapter, set status to phase3_chapters
        if paper.status == "revision":
            paper.status = "phase3_chapters"

    db.add(paper)
    db.commit()
    db.refresh(paper)
    
    if rejected_chapter:
        note = f" Chapter feedback: {comments.strip()}" if comments and comments.strip() else ""
        _notify_student(db, paper, f"Chapter {rejected_chapter} needs revision for '{paper.title}'.{note}")
    else:
        approved_count = sum(1 for value in chapter_values if value)
        note = f" Feedback: {comments.strip()}" if comments and comments.strip() else ""
        _notify_student(db, paper, f"Chapter {approved_count} was approved for '{paper.title}'.{note}")
    return _to_paper_read(paper, db, current_user)


@router.post("/papers/{paper_id}/complete-phase3", response_model=PaperRead)
def complete_phase3(
    paper_id: int,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
) -> PaperRead:
    paper = get_paper(db, paper_id)
    if not paper:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")
    
    is_supervisor = paper.supervisor_id == current_user.id or current_user.is_admin
    if not is_supervisor:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Only the assigned supervisor can complete Phase 3")
    
    paper.combined_thesis_supervisor_approved = True
    paper.status = "phase4_pending_examiners"
    db.add(paper)
    
    _record_workflow_event(
        db,
        paper_id=paper.id,
        event_type="complete_phase3",
        actor_id=current_user.id,
        actor_role="project_supervisor",
        from_status="phase3_chapters",
        to_status=paper.status,
        message=f"Phase 3 Complete: All 5 chapters approved by supervisor for '{paper.title}'. Examiners need to be assigned.",
    )
    db.commit()
    db.refresh(paper)
    
    _notify_hod_and_coordinators(
        db,
        paper,
        f"Phase 3 Complete: All 5 chapters approved by supervisor for '{paper.title}'. Examiners need to be assigned."
    )
    _notify_student(db, paper, "Congratulations! Your supervisor approved all 5 chapters. The HOD or Project Coordinator will now assign internal and external examiners.")
    
    return _to_paper_read(paper, db, current_user)


@router.post("/papers/{paper_id}/assign-examiners", response_model=PaperRead)
def assign_examiners(
    paper_id: int,
    internal_examiner_id: int = Form(...),
    external_examiner_id: int = Form(...),
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
) -> PaperRead:
    paper = get_paper(db, paper_id)
    if not paper:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")
    
    is_hod = has_role(db, current_user, "hod")
    is_coord = has_role(db, current_user, "project_coordinator")
    if not (is_hod or is_coord or current_user.is_admin):
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Only the HOD, Project Coordinator, or Admin can assign examiners")
    
    if internal_examiner_id == external_examiner_id:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Conflict of interest: Internal and external examiners must be distinct individuals",
        )

    int_exam = db.query(User).filter(User.id == internal_examiner_id).first()
    ext_exam = db.query(User).filter(User.id == external_examiner_id).first()
    if not int_exam or not ext_exam:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Selected internal or external examiner not found")
        
    int_roles = set(get_user_roles(db, internal_examiner_id))
    int_roles.add(int_exam.role)
    allowed_int = {"lecturer", "project_supervisor", "hod", "project_coordinator", "dean"}
    if not int_roles.intersection(allowed_int):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Internal examiner must be a lecturer, supervisor, HOD, or Dean")
        
    ext_roles = set(get_user_roles(db, external_examiner_id))
    ext_roles.add(ext_exam.role)
    allowed_ext = {"external_examiner", "lecturer", "project_supervisor", "hod", "project_coordinator", "dean"}
    if not ext_roles.intersection(allowed_ext):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="External examiner must be an external examiner, lecturer, HOD, or Dean")
    
    paper.internal_examiner_id = internal_examiner_id
    paper.external_examiner_id = external_examiner_id
    paper.status = "phase4_marking"
    db.add(paper)
    
    _record_workflow_event(
        db,
        paper_id=paper.id,
        event_type="assign_examiners",
        actor_id=current_user.id,
        actor_role="project_coordinator" if is_coord else "hod",
        from_status="phase4_pending_examiners",
        to_status=paper.status,
        message=f"Examiners assigned: Internal={int_exam.full_name or int_exam.email}, External={ext_exam.full_name or ext_exam.email}",
    )
    db.commit()
    db.refresh(paper)
    
    create_notification(db, user_id=internal_examiner_id, paper_id=paper.id, ntype="workflow_update", message=f"You have been assigned as Internal Examiner for '{paper.title}'. Please mark and submit results.")
    create_notification(db, user_id=external_examiner_id, paper_id=paper.id, ntype="workflow_update", message=f"You have been assigned as External Examiner for '{paper.title}'. Please mark and submit results.")
    
    _notify_student(db, paper, "Internal and external examiners have been assigned. They will now review and mark your thesis.")
    return _to_paper_read(paper, db, current_user)


@router.post("/papers/{paper_id}/upload-results", response_model=PaperRead)
async def upload_results(
    paper_id: int,
    internal_score: float | None = Form(None),
    external_score: float | None = Form(None),
    examiner_corrections: str = Form(...),
    file: UploadFile = File(None),
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
) -> PaperRead:
    paper = get_paper(db, paper_id)
    if not paper:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")
    
    is_internal = current_user.id == paper.internal_examiner_id
    is_external = current_user.id == paper.external_examiner_id
    is_hod = has_role(db, current_user, "hod")
    is_coord = has_role(db, current_user, "project_coordinator")
    is_admin = current_user.is_admin or has_role(db, current_user, "system_admin")
    
    if not (is_internal or is_external or is_hod or is_coord or is_admin):
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Only assigned examiners, Project Coordinator, HOD, or Admin can upload results")
    
    from_status = paper.status
    
    file_path = None
    file_name = None
    if file:
        safe_name = f"{uuid4().hex}_{Path(file.filename or 'results.bin').name}"
        dest = UPLOADS_DIR / safe_name
        content = await file.read()
        if content:
            dest.write_bytes(content)
            file_path = str(dest)
            file_name = file.filename

    if is_internal or is_admin or is_hod or is_coord:
        if internal_score is not None:
            paper.internal_score = internal_score
        if file_path:
            paper.internal_result_file_path = file_path
            paper.internal_result_file_name = file_name
            paper.examiner_result_file_path = file_path
            paper.examiner_result_file_name = file_name
            
    if is_external or is_admin or is_hod or is_coord:
        if external_score is not None:
            paper.external_score = external_score
        if file_path:
            paper.external_result_file_path = file_path
            paper.external_result_file_name = file_name

    role_label = "Internal Examiner" if is_internal else ("External Examiner" if is_external else "Dept Coordinator/HOD")
    if paper.examiner_corrections:
        paper.examiner_corrections = f"{paper.examiner_corrections}\n\n[{role_label} - {current_user.full_name or current_user.email}]: {examiner_corrections}"
    else:
        paper.examiner_corrections = f"[{role_label} - {current_user.full_name or current_user.email}]: {examiner_corrections}"

    both_marked = paper.internal_score is not None and paper.external_score is not None
    if both_marked:
        paper.status = "phase5_corrections"
    else:
        paper.status = "phase4_marking"

    db.add(paper)
    
    _record_workflow_event(
        db,
        paper_id=paper.id,
        event_type="upload_results",
        actor_id=current_user.id,
        actor_role="examiner" if (is_internal or is_external) else "project_coordinator",
        from_status=from_status,
        to_status=paper.status,
        message=f"Marking update submitted by {current_user.full_name or current_user.email}.",
    )
    db.commit()
    db.refresh(paper)
    
    if both_marked:
        _notify_hod_and_coordinators(
            db,
            paper,
            f"Phase 4 Complete: Both examiner marking results and corrections uploaded for '{paper.title}'."
        )
        _notify_student(db, paper, "Examiners have completed their markings. Please check the examiner corrections, make the necessary adjustments, and upload the updated document.")
    else:
        _notify_hod_and_coordinators(
            db,
            paper,
            f"Examiner marking updated: Partial score received for '{paper.title}' from {current_user.full_name or current_user.email}."
        )

    return _to_paper_read(paper, db, current_user)


@router.post("/papers/{paper_id}/upload-corrections", response_model=PaperRead)
async def upload_corrections(
    paper_id: int,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
) -> PaperRead:
    paper = get_paper(db, paper_id)
    if not paper:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")
    
    if paper.created_by_id != current_user.id and not current_user.is_admin:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Only the student author can upload final corrections")
    
    safe_name = f"{uuid4().hex}_{Path(file.filename or 'upload.bin').name}"
    dest = UPLOADS_DIR / safe_name
    content = await file.read()
    if not content:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Uploaded file is empty")
    dest.write_bytes(content)
    
    paper.file_name = file.filename
    paper.file_path = str(dest)
    paper.file_size = len(content)
    paper.mime_type = file.content_type
    paper.status = "phase5_pending_supervisor"
    db.add(paper)
    
    _record_workflow_event(
        db,
        paper_id=paper.id,
        event_type="upload_corrections",
        actor_id=current_user.id,
        actor_role="student",
        from_status="phase5_corrections",
        to_status=paper.status,
        message="Student uploaded corrections, awaiting supervisor approval.",
    )
    db.commit()
    db.refresh(paper)
    
    if paper.supervisor_id:
        create_notification(
            db,
            user_id=paper.supervisor_id,
            paper_id=paper.id,
            ntype="workflow_update",
            message=f"Student uploaded final corrections for '{paper.title}'. Please review and approve."
        )
        
    return _to_paper_read(paper, db, current_user)


@router.post("/papers/{paper_id}/submit-in-system-corrections", response_model=PaperRead)
async def submit_in_system_corrections(
    paper_id: int,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
) -> PaperRead:
    paper = get_paper(db, paper_id)
    if not paper:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")
    
    if paper.created_by_id != current_user.id and not current_user.is_admin:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Only the student author can submit corrections")
    
    if not paper.file_path or not Path(paper.file_path).exists():
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="No in-system document found to submit")
    
    paper.status = "phase5_pending_supervisor"
    db.add(paper)
    
    _record_workflow_event(
        db,
        paper_id=paper.id,
        event_type="submit_in_system_corrections",
        actor_id=current_user.id,
        actor_role="student",
        from_status="phase5_corrections",
        to_status=paper.status,
        message="Student submitted in-system ONLYOFFICE edits as final corrections, awaiting supervisor approval.",
    )
    db.commit()
    db.refresh(paper)
    
    if paper.supervisor_id:
        create_notification(
            db,
            user_id=paper.supervisor_id,
            paper_id=paper.id,
            ntype="workflow_update",
            message=f"Student submitted in-system ONLYOFFICE corrections for '{paper.title}'. Please review and approve."
        )
        
    return _to_paper_read(paper, db, current_user)


@router.post("/papers/{paper_id}/supervisor-approve-corrections", response_model=PaperRead)
def supervisor_approve_corrections(
    paper_id: int,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
) -> PaperRead:
    paper = get_paper(db, paper_id)
    if not paper:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")
    
    is_supervisor = paper.supervisor_id == current_user.id or current_user.is_admin or has_role(db, current_user, "hod") or has_role(db, current_user, "project_coordinator")
    if not is_supervisor:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Only the assigned supervisor, HOD, or coordinator can approve corrections")
    
    paper.status = "phase5_pending_coordinator"
    db.add(paper)
    
    _record_workflow_event(
        db,
        paper_id=paper.id,
        event_type="supervisor_approve_corrections",
        actor_id=current_user.id,
        actor_role="project_supervisor",
        from_status="phase5_pending_supervisor",
        to_status=paper.status,
        message="Supervisor approved corrections, awaiting coordinator sign-off.",
    )
    db.commit()
    db.refresh(paper)
    
    _notify_hod_and_coordinators(
        db,
        paper,
        f"Supervisor approved corrections for '{paper.title}', awaiting final sign-off."
    )
    
    return _to_paper_read(paper, db, current_user)


@router.post("/papers/{paper_id}/supervisor-reject-corrections", response_model=PaperRead)
def supervisor_reject_corrections(
    paper_id: int,
    feedback: str = Form("Corrections not satisfactorily addressed. Please review examiner comments and make required revisions."),
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
) -> PaperRead:
    paper = get_paper(db, paper_id)
    if not paper:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")
    
    is_supervisor = paper.supervisor_id == current_user.id or current_user.is_admin or has_role(db, current_user, "hod") or has_role(db, current_user, "project_coordinator")
    if not is_supervisor:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Only the assigned supervisor or HOD can request revisions")
    
    paper.status = "phase5_corrections"
    if feedback.strip():
        paper.review_comments = feedback.strip()
    db.add(paper)
    db.commit()
    db.refresh(paper)
    
    # Notify student
    if paper.created_by_id:
        from app.models.thesis_system import Notification
        notif = Notification(
            user_id=paper.created_by_id,
            title="Further Corrections Required",
            message=f"Supervisor reviewed your resubmitted work for '{paper.title}' and requested further corrections: {feedback.strip()}",
            ntype="workflow_update",
        )
        db.add(notif)
        db.commit()
        
    return _to_paper_read(paper, db, current_user)


@router.post("/papers/{paper_id}/coordinator-approve-corrections", response_model=PaperRead)
def coordinator_approve_corrections(
    paper_id: int,
    decision: str = Form("approved"),
    comment: str | None = Form(None),
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
) -> PaperRead:
    paper = get_paper(db, paper_id)
    if not paper:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")

    allowed = (
        current_user.is_admin
        or has_role(db, current_user, "project_coordinator")
        or has_role(db, current_user, "hod")
        or has_role(db, current_user, "system_admin")
    )
    if not allowed:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Only Project Coordinator or Admin can approve corrections")

    if decision == "approved":
        paper.project_coordinator_approved_by_id = current_user.id
        paper.project_coordinator_approved_at = datetime.now(timezone.utc)
        
        if paper.hod_approved_at is not None or has_role(db, current_user, "hod") or current_user.is_admin:
            paper.status = "phase5_approved_for_library"
        else:
            paper.status = "phase5_pending_hod"
        
        _record_workflow_event(
            db,
            paper_id=paper.id,
            event_type="coordinator_approve_corrections",
            actor_id=current_user.id,
            actor_role="project_coordinator",
            from_status="phase5_pending_coordinator",
            to_status=paper.status,
            message="Project Coordinator approved corrections.",
        )
    else:
        paper.status = "phase5_corrections"
        paper.review_comments = comment or "Revisions requested by Project Coordinator."
        _record_workflow_event(
            db,
            paper_id=paper.id,
            event_type="coordinator_reject_corrections",
            actor_id=current_user.id,
            actor_role="project_coordinator",
            from_status="phase5_pending_coordinator",
            to_status=paper.status,
            message=f"Project Coordinator requested revisions: {comment}",
        )

    db.commit()
    db.refresh(paper)

    if paper.created_by_id:
        notif_msg = f"Project Coordinator reviewed your corrections: {decision.upper()}."
        create_notification(
            db,
            user_id=paper.created_by_id,
            paper_id=paper.id,
            ntype="workflow_update",
            message=notif_msg,
        )

    return _to_paper_read(paper, db, current_user)


@router.post("/papers/{paper_id}/hod-approve-corrections", response_model=PaperRead)
def hod_approve_corrections(
    paper_id: int,
    decision: str = Form("approved"),
    comment: str | None = Form(None),
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
) -> PaperRead:
    paper = get_paper(db, paper_id)
    if not paper:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")

    allowed = (
        current_user.is_admin
        or has_role(db, current_user, "hod")
        or has_role(db, current_user, "system_admin")
    )
    if not allowed:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Only Head of Department (HOD) or Admin can approve corrections")

    if decision == "approved":
        paper.hod_approved_by_id = current_user.id
        paper.hod_approved_at = datetime.now(timezone.utc)
        
        if paper.project_coordinator_approved_at is not None or current_user.is_admin:
            paper.status = "phase5_approved_for_library"
        else:
            paper.status = "phase5_pending_coordinator"

        _record_workflow_event(
            db,
            paper_id=paper.id,
            event_type="hod_approve_corrections",
            actor_id=current_user.id,
            actor_role="hod",
            from_status="phase5_pending_hod",
            to_status=paper.status,
            message="HOD signed and approved corrections.",
        )
    else:
        paper.status = "phase5_corrections"
        paper.review_comments = comment or "Revisions requested by HOD."
        _record_workflow_event(
            db,
            paper_id=paper.id,
            event_type="hod_reject_corrections",
            actor_id=current_user.id,
            actor_role="hod",
            from_status="phase5_pending_hod",
            to_status=paper.status,
            message=f"HOD requested revisions: {comment}",
        )

    db.commit()
    db.refresh(paper)

    if paper.created_by_id:
        notif_msg = f"Head of Department (HOD) reviewed your corrections: {decision.upper()}."
        create_notification(
            db,
            user_id=paper.created_by_id,
            paper_id=paper.id,
            ntype="workflow_update",
            message=notif_msg,
        )

    return _to_paper_read(paper, db, current_user)


@router.get("/papers/{paper_id}/examiner-script/internal")
def download_internal_examiner_script(
    paper_id: int,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    paper = get_paper(db, paper_id)
    if not paper:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")
    
    is_author = paper.created_by_id == current_user.id
    is_supervisor = paper.supervisor_id == current_user.id
    is_admin = current_user.is_admin or has_role(db, current_user, "system_admin")
    
    is_hod_or_coord = False
    user_dept = (current_user.department or "").strip().lower()
    paper_dept = ""
    if paper.department:
        paper_dept = (paper.department.name or "").strip().lower()
    elif paper.discipline:
        paper_dept = (paper.discipline or "").strip().lower()
        
    if user_dept and paper_dept and user_dept == paper_dept:
        if has_role(db, current_user, "hod") or has_role(db, current_user, "project_coordinator"):
            is_hod_or_coord = True
            
    if not (is_author or is_supervisor or is_admin or is_hod_or_coord):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="You are not authorized to download this examiner script"
        )
        
    if not paper.internal_result_file_path:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Internal examiner script not uploaded"
        )
        
    path = Path(paper.internal_result_file_path)
    if not path.exists() or not path.is_file():
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Internal examiner script file not found on server"
        )
        
    return FileResponse(
        path=path,
        media_type="application/octet-stream",
        filename=paper.internal_result_file_name or path.name,
    )


@router.get("/papers/{paper_id}/examiner-script/external")
def download_external_examiner_script(
    paper_id: int,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    paper = get_paper(db, paper_id)
    if not paper:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")
    
    is_author = paper.created_by_id == current_user.id
    is_supervisor = paper.supervisor_id == current_user.id
    is_admin = current_user.is_admin or has_role(db, current_user, "system_admin")
    
    is_hod_or_coord = False
    user_dept = (current_user.department or "").strip().lower()
    paper_dept = ""
    if paper.department:
        paper_dept = (paper.department.name or "").strip().lower()
    elif paper.discipline:
        paper_dept = (paper.discipline or "").strip().lower()
        
    if user_dept and paper_dept and user_dept == paper_dept:
        if has_role(db, current_user, "hod") or has_role(db, current_user, "project_coordinator"):
            is_hod_or_coord = True
            
    if not (is_author or is_supervisor or is_admin or is_hod_or_coord):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="You are not authorized to download this examiner script"
        )
        
    if not paper.external_result_file_path:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="External examiner script not uploaded"
        )
        
    path = Path(paper.external_result_file_path)
    if not path.exists() or not path.is_file():
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="External examiner script file not found on server"
        )
        
    return FileResponse(
        path=path,
        media_type="application/octet-stream",
        filename=paper.external_result_file_name or path.name,
    )


@router.post("/papers/{paper_id}/upload-combined-thesis", response_model=PaperRead)
async def upload_combined_thesis(
    paper_id: int,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
) -> PaperRead:
    paper = get_paper(db, paper_id)
    if not paper:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")
    
    if paper.created_by_id != current_user.id and not current_user.is_admin:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Only the student author can upload the combined thesis")
    
    if not paper.ch5_supervisor_approved:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Cannot upload combined thesis: Chapter 5 must be approved first")
    
    content = await file.read()
    if not content:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Uploaded file is empty")
    
    safe_name = f"{uuid4().hex}_{Path(file.filename or 'combined.bin').name}"
    dest = UPLOADS_DIR / safe_name
    dest.write_bytes(content)
    
    paper.file_name = file.filename
    paper.file_path = str(dest)
    paper.file_size = len(content)
    paper.mime_type = file.content_type
    paper.combined_thesis_student_done = True
    paper.combined_thesis_supervisor_approved = False
    
    db.add(paper)
    
    _record_paper_version(
        db,
        paper=paper,
        source="student_combined_thesis",
        actor_id=current_user.id,
        note="Student uploaded combined thesis document",
    )
    _record_workflow_event(
        db,
        paper_id=paper.id,
        event_type="upload_combined_thesis",
        actor_id=current_user.id,
        actor_role="student",
        from_status=paper.status,
        to_status=paper.status,
        message="Combined thesis uploaded by student, awaiting supervisor sign-off.",
    )
    db.commit()
    db.refresh(paper)
    
    if paper.supervisor_id:
        create_notification(
            db,
            user_id=paper.supervisor_id,
            paper_id=paper.id,
            ntype="workflow_update",
            message=f"Student uploaded combined thesis for '{paper.title}'. Please review and sign off."
        )
    
    return _to_paper_read(paper, db, current_user)


@router.post("/papers/{paper_id}/combined-thesis/supervisor", response_model=PaperRead)
def supervisor_approve_combined_thesis(
    paper_id: int,
    approved: bool = Form(...),
    comments: str | None = Form(None),
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
) -> PaperRead:
    paper = get_paper(db, paper_id)
    if not paper:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")
    
    is_supervisor = paper.supervisor_id == current_user.id or current_user.is_admin
    if not is_supervisor:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Only the assigned supervisor or admin can approve the combined thesis")
    
    if not paper.combined_thesis_student_done:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Combined thesis has not been uploaded by the student yet")

    if approved:
        paper.combined_thesis_supervisor_approved = True
        from_status = paper.status
        paper.status = "phase4_pending_examiners"
        db.add(paper)
        
        msg = f"Supervisor signed off and approved the combined thesis for '{paper.title}'."
        if comments and comments.strip():
            msg += f" Comments: {comments.strip()}"
            
        _record_workflow_event(
            db,
            paper_id=paper.id,
            event_type="approve_combined_thesis",
            actor_id=current_user.id,
            actor_role="project_supervisor",
            from_status=from_status,
            to_status=paper.status,
            message=msg,
        )
        db.commit()
        db.refresh(paper)
        
        _notify_hod_and_coordinators(
            db,
            paper,
            f"Combined thesis approved by supervisor for '{paper.title}'. Examiners need to be assigned."
        )
        _notify_student(
            db,
            paper,
            "Your supervisor has approved your combined thesis! The paper is now entering Phase 4 (Examiner Assignment)."
        )
    else:
        paper.combined_thesis_student_done = False
        paper.combined_thesis_supervisor_approved = False
        from_status = paper.status
        paper.status = "revision"
        db.add(paper)
        
        msg = f"Supervisor requested revision for the combined thesis of '{paper.title}'."
        if comments and comments.strip():
            msg += f" Feedback: {comments.strip()}"
            
        _record_workflow_event(
            db,
            paper_id=paper.id,
            event_type="reject_combined_thesis",
            actor_id=current_user.id,
            actor_role="project_supervisor",
            from_status=from_status,
            to_status=paper.status,
            message=msg,
        )
        db.commit()
        db.refresh(paper)
        
        feedback = f" Feedback comments: {comments.strip()}" if comments and comments.strip() else ""
        _notify_student(
            db,
            paper,
            f"Your combined thesis for '{paper.title}' needs revision.{feedback}"
        )
        
    return _to_paper_read(paper, db, current_user)


@router.post("/papers/{paper_id}/upload-draft", response_model=PaperRead)
async def upload_draft(
    paper_id: int,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
) -> PaperRead:
    paper = get_paper(db, paper_id)
    if not paper:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Paper not found")
    
    if paper.created_by_id != current_user.id and not current_user.is_admin:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Only the student author can upload drafts")
    
    if paper.status not in {"phase3_chapters", "revision"}:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Can only upload drafts during the chapters writing phase (Phase 3)")
    
    content = await file.read()
    if not content:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Uploaded file is empty")
    
    safe_name = f"{uuid4().hex}_{Path(file.filename or 'draft.bin').name}"
    dest = UPLOADS_DIR / safe_name
    dest.write_bytes(content)
    
    paper.file_name = file.filename
    paper.file_path = str(dest)
    paper.file_size = len(content)
    paper.mime_type = file.content_type
    
    # Automatically set the student done checklist item to True for the current active chapter
    if not paper.ch1_supervisor_approved:
        paper.ch1_student_done = True
    elif not paper.ch2_supervisor_approved:
        paper.ch2_student_done = True
    elif not paper.ch3_supervisor_approved:
        paper.ch3_student_done = True
    elif not paper.ch4_supervisor_approved:
        paper.ch4_student_done = True
    elif not paper.ch5_supervisor_approved:
        paper.ch5_student_done = True

    # If the paper was in revision status, change it back to active chapter reviews phase
    if paper.status == "revision":
        paper.status = "phase3_chapters"

    db.add(paper)
    
    _record_paper_version(
        db,
        paper=paper,
        source="student_draft_upload",
        actor_id=current_user.id,
        note=f"Student uploaded new draft: {file.filename}",
    )
    _record_workflow_event(
        db,
        paper_id=paper.id,
        event_type="upload_draft",
        actor_id=current_user.id,
        actor_role="student",
        from_status=paper.status,
        to_status=paper.status,
        message=f"New draft uploaded by student: {file.filename}",
    )
    db.commit()
    db.refresh(paper)
    
    if paper.supervisor_id:
        create_notification(
            db,
            user_id=paper.supervisor_id,
            paper_id=paper.id,
            ntype="workflow_update",
            message=f"Student uploaded a new draft of their thesis: {file.filename}"
        )
    
    return _to_paper_read(paper, db, current_user)


@router.get("/papers/examiner/download-zip")
def download_examiner_assigned_zip(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Bulk download all papers assigned to the examiner in a zipped archive."""
    import io
    import zipfile
    from pathlib import Path
    from fastapi.responses import StreamingResponse
    
    is_examiner = (
        db.query(Paper)
        .filter(
            (Paper.status == "phase4_marking")
            & (
                (Paper.internal_examiner_id == current_user.id)
                | (Paper.external_examiner_id == current_user.id)
            )
        )
        .first()
    )
    if not is_examiner and not current_user.is_admin:
         raise HTTPException(
             status_code=status.HTTP_403_FORBIDDEN,
             detail="Only assigned examiners or admin can download this ZIP"
         )
         
    papers = (
        db.query(Paper)
        .filter(
            (Paper.status == "phase4_marking")
            & (
                (Paper.internal_examiner_id == current_user.id)
                | (Paper.external_examiner_id == current_user.id)
            )
        )
        .all()
    )
    
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        added_any = False
        added_filenames = set()
        for paper in papers:
            if not paper.file_path:
                continue
            path = Path(paper.file_path)
            if not path.exists() or not path.is_file():
                continue
                
            student_id = paper.created_by.school_id if (paper.created_by and paper.created_by.school_id) else f"student-{paper.created_by_id or 'unknown'}"
            student_name = paper.created_by.full_name if (paper.created_by and paper.created_by.full_name) else "Unknown"
            
            title_slug = "".join(c if c.isalnum() or c in (" ", "_", "-") else "" for c in paper.title)
            title_slug = title_slug.replace(" ", "_")[:50]
            suffix = path.suffix or ".pdf"
            filename = f"{student_id}_{student_name}_{title_slug}{suffix}"
            
            counter = 1
            while filename in added_filenames:
                filename = f"{student_id}_{student_name}_{title_slug}_{counter}{suffix}"
                counter += 1
            added_filenames.add(filename)
            
            zip_file.write(path, arcname=filename)
            added_any = True
            
        if not added_any:
            zip_file.writestr("README.txt", "No papers assigned to you for marking were found.")
            
    zip_buffer.seek(0)
    headers = {
        "Content-Disposition": f'attachment; filename="examiner_{current_user.id}_assigned_works.zip"'
    }
    return StreamingResponse(zip_buffer, media_type="application/zip", headers=headers)


# -----------------------------------------------------------------------------
# Section 6 Specification REST Endpoint Aliases
# -----------------------------------------------------------------------------

@router.post("/theses/{paper_id}/topic/accept", response_model=PaperRead)
def topic_accept_alias(
    paper_id: int,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
) -> PaperRead:
    """Canonical spec alias: HOD accepts student topic."""
    paper = get_paper(db, paper_id)
    if not paper:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Thesis not found")
    paper.status = "topic_accepted"
    db.add(paper)
    db.commit()
    db.refresh(paper)
    return _to_paper_read(paper, db, current_user)


@router.post("/theses/{paper_id}/assign-supervisor", response_model=PaperRead)
def thesis_assign_supervisor_alias(
    paper_id: int,
    supervisor_id: int = Form(...),
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
) -> PaperRead:
    """Canonical spec alias: HOD assigns supervisor."""
    return assign_supervisor(paper_id=paper_id, supervisor_id=supervisor_id, db=db, current_user=current_user)







@router.post("/theses/{paper_id}/finish-steps", response_model=PaperRead)
def thesis_finish_steps_alias(
    paper_id: int,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
) -> PaperRead:
    """Canonical spec alias: Supervisor finishes steps and advances thesis to Phase 3 examination."""
    return complete_phase3(paper_id=paper_id, db=db, current_user=current_user)


@router.post("/theses/{paper_id}/assign-examiners", response_model=PaperRead)
def thesis_assign_examiners_alias(
    paper_id: int,
    internal_examiner_id: int = Form(...),
    external_examiner_id: int = Form(...),
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
) -> PaperRead:
    """Canonical spec alias: HOD assigns internal and external examiners."""
    return assign_examiners(
        paper_id=paper_id,
        internal_examiner_id=internal_examiner_id,
        external_examiner_id=external_examiner_id,
        db=db,
        current_user=current_user,
    )


@router.get("/examiners/{examiner_id}/download-zip")
def examiner_download_zip_alias(
    examiner_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Canonical spec alias: Examiner downloads assigned student ZIP archive."""
    return download_examiner_assigned_zip(db=db, current_user=current_user)


@router.post("/examiner-assignments/{assignment_id}/upload-marks", response_model=PaperRead)
async def examiner_upload_marks_alias(
    assignment_id: int,
    internal_score: float | None = Form(None),
    external_score: float | None = Form(None),
    comments: str = Form(...),
    file: UploadFile = File(None),
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
) -> PaperRead:
    """Canonical spec alias: Examiner uploads marking results."""
    return await upload_results(
        paper_id=assignment_id,
        internal_score=internal_score,
        external_score=external_score,
        examiner_corrections=comments,
        file=file,
        db=db,
        current_user=current_user,
    )


@router.post("/theses/{paper_id}/compile-comments", response_model=PaperRead)
def thesis_compile_comments_alias(
    paper_id: int,
    compiled_comments: str = Form(...),
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
) -> PaperRead:
    """Canonical spec alias: HOD compiles examiner comments."""
    payload = PaperReview(decision="revision", comments=compiled_comments)
    return review_paper_endpoint(paper_id=paper_id, payload=payload, db=db, current_admin=current_user)


@router.post("/theses/{paper_id}/send-comments-to-student", response_model=PaperRead)
def thesis_send_comments_alias(
    paper_id: int,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
) -> PaperRead:
    """Canonical spec alias: HOD relays comments to student advancing to Phase 4."""
    paper = get_paper(db, paper_id)
    if not paper:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Thesis not found")
    paper.status = "phase5_pending_supervisor"
    db.add(paper)
    db.commit()
    db.refresh(paper)
    _notify_student(db, paper, "Examiner comments compiled. Please review comments and submit corrections.")
    return _to_paper_read(paper, db, current_user)


@router.post("/theses/{paper_id}/corrections", response_model=PaperRead)
def thesis_submit_corrections_alias(
    paper_id: int,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
) -> PaperRead:
    """Canonical spec alias: Student submits Phase 4 corrections."""
    return upload_corrections(paper_id=paper_id, file=file, db=db, current_user=current_user)


@router.post("/corrections/{correction_id}/supervisor-decision", response_model=PaperRead)
def correction_supervisor_decision_alias(
    correction_id: int,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
) -> PaperRead:
    """Canonical spec alias: Supervisor approves Phase 4 corrections."""
    return supervisor_approve_corrections(paper_id=correction_id, db=db, current_user=current_user)


@router.post("/corrections/{correction_id}/hod-decision", response_model=PaperRead)
def correction_hod_decision_alias(
    correction_id: int,
    comments: str | None = Form(None),
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
) -> PaperRead:
    """Canonical spec alias: HOD approves Phase 4 corrections (dual sign-off)."""
    payload = PaperReview(decision="approve", comments=comments)
    return review_paper_endpoint(paper_id=correction_id, payload=payload, db=db, current_admin=current_user)


@router.post("/theses/{paper_id}/publish", response_model=PaperRead)
def thesis_publish_alias(
    paper_id: int,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
) -> PaperRead:
    """Canonical spec alias: Librarian publishes thesis to repository."""
    payload = PaperReview(decision="approve", comments="Published by Librarian")
    return review_paper_endpoint(paper_id=paper_id, payload=payload, db=db, current_admin=current_user)


@router.get("/repository", response_model=list[PaperRead])
def repository_browse_alias(
    department: str | None = Query(None),
    year: int | None = Query(None),
    search: str | None = Query(None),
    db: Session = Depends(get_db),
) -> list[PaperRead]:
    """Canonical spec alias: Public repository browse page."""
    return list_papers(db=db, discipline=department, year=year, q=search, catalog_mode=True)


@router.get("/repository/{paper_id}", response_model=PaperRead)
def repository_get_alias(
    paper_id: int,
    db: Session = Depends(get_db),
) -> PaperRead:
    """Canonical spec alias: Read published thesis detail."""
    paper = get_paper(db, paper_id)
    if not paper or not paper.is_public:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Published thesis not found")
    return _to_paper_read(paper, db)


@router.get("/theses/{paper_id}/comments")
def thesis_get_comments_alias(
    paper_id: int,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    """Canonical spec alias: External comment channel list."""
    return get_annotations_endpoint(paper_id=paper_id, db=db, current_user=current_user)


@router.post("/theses/{paper_id}/comments")
def thesis_add_comment_alias(
    paper_id: int,
    text: str = Form(...),
    location: str | None = Form(None),
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    """Canonical spec alias: Add comment to external comment panel channel."""
    return create_annotation_endpoint(paper_id=paper_id, text=text, location=location, db=db, current_user=current_user)


@router.get("/hod/dashboard")
def hod_dashboard_alias(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Canonical spec alias: HOD / Project Coordinator dashboard overview."""
    return read_paper_stats(db=db)


@router.get("/dean/dashboard")
def dean_dashboard_alias(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Canonical spec alias: Dean school-wide dashboard rollup overview."""
    return read_paper_stats(db=db)


@router.get("/papers/reports/export")
def export_academic_report(
    degree_level: str | None = Query(None),
    department: str | None = Query(None),
    lecturer_id: int | None = Query(None),
    student_id: int | None = Query(None),
    status_filter: str | None = Query(None),
    format: str = Query("xlsx"),
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    is_admin_or_dean = current_user.is_admin or has_role(db, current_user, "system_admin") or has_role(db, current_user, "dean") or has_role(db, current_user, "head_library") or has_role(db, current_user, "librarian")
    is_hod_or_coord = has_role(db, current_user, "hod") or has_role(db, current_user, "project_coordinator")
    is_lecturer = has_role(db, current_user, "lecturer") or current_user.role == "lecturer"
    
    if not (is_admin_or_dean or is_hod_or_coord or is_lecturer):
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Access denied. Only academic staff and administrators can generate reports.")
        
    query = db.query(Paper)
    
    if not is_admin_or_dean:
        if is_hod_or_coord:
            u_dept = (current_user.department or "").strip().lower()
            if u_dept:
                query = query.filter(
                    (func.lower(func.coalesce(Paper.discipline, "")) == u_dept)
                    | (Paper.department_id != None)
                )
        elif is_lecturer:
            query = query.filter(
                (Paper.supervisor_id == current_user.id)
                | (Paper.internal_examiner_id == current_user.id)
                | (Paper.external_examiner_id == current_user.id)
            )
            
    if department and department.strip().lower() not in {"all", ""}:
        d_val = department.strip().lower()
        query = query.filter(
            (func.lower(func.coalesce(Paper.discipline, "")).contains(d_val))
            | (func.lower(func.coalesce(Paper.university, "")).contains(d_val))
        )
        
    if lecturer_id and lecturer_id > 0:
        query = query.filter(
            (Paper.supervisor_id == lecturer_id)
            | (Paper.internal_examiner_id == lecturer_id)
            | (Paper.external_examiner_id == lecturer_id)
        )
        
    if student_id and student_id > 0:
        query = query.filter(Paper.created_by_id == student_id)
        
    if status_filter and status_filter.strip().lower() not in {"all", ""}:
        query = query.filter(Paper.status == status_filter.strip())
        
    papers = query.order_by(Paper.id.desc()).all()
    
    if degree_level and degree_level.strip().lower() not in {"all", ""}:
        target_deg = degree_level.strip().lower()
        filtered_papers = []
        for p in papers:
            stu = db.query(User).filter(User.id == p.created_by_id).first() if p.created_by_id else None
            p_deg = classify_degree_level(paper=p, student_user=stu, db=db).lower()
            if target_deg in p_deg or p_deg in target_deg:
                filtered_papers.append(p)
        papers = filtered_papers

    report_rows = []
    for idx, p in enumerate(papers, start=1):
        stu = db.query(User).filter(User.id == p.created_by_id).first() if p.created_by_id else None
        deg = classify_degree_level(paper=p, student_user=stu, db=db)
        
        sup_user = db.query(User).filter(User.id == p.supervisor_id).first() if p.supervisor_id else None
        int_user = db.query(User).filter(User.id == p.internal_examiner_id).first() if p.internal_examiner_id else None
        ext_user = db.query(User).filter(User.id == p.external_examiner_id).first() if p.external_examiner_id else None
        
        int_m = p.internal_score
        ext_m = p.external_score
        avg_m = None
        if int_m is not None and ext_m is not None:
            avg_m = round((int_m + ext_m) / 2.0, 1)
        elif int_m is not None:
            avg_m = int_m
        elif ext_m is not None:
            avg_m = ext_m
            
        grade = "N/A"
        if avg_m is not None:
            if avg_m >= 80: grade = "A (Distinction)"
            elif avg_m >= 75: grade = "B+ (Very Good)"
            elif avg_m >= 70: grade = "B (Good)"
            elif avg_m >= 65: grade = "C+ (Credit)"
            elif avg_m >= 60: grade = "C (Pass)"
            elif avg_m >= 55: grade = "D+ (Marginal Pass)"
            elif avg_m >= 50: grade = "D (Pass)"
            else: grade = "F (Fail)"
            
        report_rows.append({
            "num": idx,
            "id": f"PAPER-{p.id}",
            "student_id": stu.school_id if (stu and getattr(stu, 'school_id', None)) else (f"STU-{stu.id}" if stu else "-"),
            "student_name": stu.full_name if stu else (p.authors[0].name if p.authors else "Unknown"),
            "degree_level": deg,
            "discipline": p.discipline or (stu.program if stu else "-") or "Computer Science",
            "title": p.title,
            "supervisor": sup_user.full_name if sup_user else "Unassigned",
            "internal_examiner": int_user.full_name if int_user else "-",
            "external_examiner": ext_user.full_name if ext_user else "-",
            "internal_mark": int_m if int_m is not None else "-",
            "external_mark": ext_m if ext_m is not None else "-",
            "final_mark": avg_m if avg_m is not None else "-",
            "grade": grade,
            "status": p.status.replace("_", " ").title(),
            "created_at": p.created_at.strftime("%Y-%m-%d") if p.created_at else "-",
        })

    if format == "csv":
        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow([
            "#", "Paper ID", "Student Index / ID", "Student Name", "Degree Level", 
            "Program / Discipline", "Thesis Title", "Supervisor", "Internal Examiner", 
            "External Examiner", "Internal Mark", "External Mark", "Final Average Mark", 
            "Letter Grade", "Status", "Date Submitted"
        ])
        for r in report_rows:
            writer.writerow([
                r["num"], r["id"], r["student_id"], r["student_name"], r["degree_level"],
                r["discipline"], r["title"], r["supervisor"], r["internal_examiner"],
                r["external_examiner"], r["internal_mark"], r["external_mark"], r["final_mark"],
                r["grade"], r["status"], r["created_at"]
            ])
        output.seek(0)
        return StreamingResponse(
            io.BytesIO(output.getvalue().encode("utf-8-sig")),
            media_type="text/csv",
            headers={"Content-Disposition": f'attachment; filename="GIMPA_Academic_Report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.csv"'}
        )
        
    elif format == "json":
        return report_rows
        
    else:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Academic Evaluation Report"
        ws.views.sheetView[0].showGridLines = True
        
        title_font = Font(name="Calibri", size=14, bold=True, color="FFFFFF")
        title_fill = PatternFill(start_color="1E1B4B", end_color="1E1B4B", fill_type="solid")
        meta_font = Font(name="Calibri", size=10, italic=True, color="475569")
        header_font = Font(name="Calibri", size=10, bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="3B82F6", end_color="3B82F6", fill_type="solid")
        data_font = Font(name="Calibri", size=10)
        bold_font = Font(name="Calibri", size=10, bold=True)
        thin_border = Border(
            left=Side(style='thin', color='CBD5E1'),
            right=Side(style='thin', color='CBD5E1'),
            top=Side(style='thin', color='CBD5E1'),
            bottom=Side(style='thin', color='CBD5E1')
        )
        
        ws.merge_cells("A1:P1")
        top_cell = ws["A1"]
        top_cell.value = "GHANA INSTITUTE OF MANAGEMENT AND PUBLIC ADMINISTRATION (GIMPA)"
        top_cell.font = title_font
        top_cell.fill = title_fill
        top_cell.alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[1].height = 30
        
        ws.merge_cells("A2:P2")
        sub_cell = ws["A2"]
        sub_cell.value = "Thesis & Academic Project Evaluation Master Report"
        sub_cell.font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
        sub_cell.fill = PatternFill(start_color="312E81", end_color="312E81", fill_type="solid")
        sub_cell.alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[2].height = 22
        
        ws.cell(row=3, column=1, value=f"Generated By: {current_user.full_name or current_user.email} | Export Date: {datetime.now().strftime('%d-%b-%Y %H:%M:%S')} | Total Candidates: {len(report_rows)}").font = meta_font
        ws.row_dimensions[3].height = 18
        
        headers = [
            "#", "Paper ID", "Student ID", "Candidate Name", "Degree Level", 
            "Program / Discipline", "Thesis Title", "Supervisor", "Internal Examiner", 
            "External Examiner", "Internal (/100)", "External (/100)", "Final Score (/100)", 
            "Grade", "Workflow Status", "Submitted Date"
        ]
        
        header_row = 4
        ws.row_dimensions[header_row].height = 24
        for col_idx, h_text in enumerate(headers, start=1):
            c = ws.cell(row=header_row, column=col_idx, value=h_text)
            c.font = header_font
            c.fill = header_fill
            c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            c.border = thin_border
            
        cur_row = 5
        for r in report_rows:
            ws.cell(row=cur_row, column=1, value=r["num"]).alignment = Alignment(horizontal="center")
            ws.cell(row=cur_row, column=2, value=r["id"]).alignment = Alignment(horizontal="center")
            ws.cell(row=cur_row, column=3, value=r["student_id"]).alignment = Alignment(horizontal="center")
            ws.cell(row=cur_row, column=4, value=r["student_name"]).font = bold_font
            ws.cell(row=cur_row, column=5, value=r["degree_level"]).alignment = Alignment(horizontal="center")
            ws.cell(row=cur_row, column=6, value=r["discipline"])
            ws.cell(row=cur_row, column=7, value=r["title"])
            ws.cell(row=cur_row, column=8, value=r["supervisor"])
            ws.cell(row=cur_row, column=9, value=r["internal_examiner"])
            ws.cell(row=cur_row, column=10, value=r["external_examiner"])
            
            c_int = ws.cell(row=cur_row, column=11, value=r["internal_mark"])
            c_int.alignment = Alignment(horizontal="center")
            c_ext = ws.cell(row=cur_row, column=12, value=r["external_mark"])
            c_ext.alignment = Alignment(horizontal="center")
            c_fin = ws.cell(row=cur_row, column=13, value=r["final_mark"])
            c_fin.alignment = Alignment(horizontal="center")
            c_fin.font = bold_font
            
            c_grd = ws.cell(row=cur_row, column=14, value=r["grade"])
            c_grd.alignment = Alignment(horizontal="center")
            c_grd.font = bold_font
            
            ws.cell(row=cur_row, column=15, value=r["status"]).alignment = Alignment(horizontal="center")
            ws.cell(row=cur_row, column=16, value=r["created_at"]).alignment = Alignment(horizontal="center")
            
            for col_i in range(1, 17):
                cell = ws.cell(row=cur_row, column=col_i)
                cell.border = thin_border
                if not cell.font.bold:
                    cell.font = data_font
            ws.row_dimensions[cur_row].height = 20
            cur_row += 1
            
        ws.row_dimensions[cur_row].height = 22
        ws.cell(row=cur_row, column=1, value="TOTALS").font = bold_font
        ws.cell(row=cur_row, column=2, value=f"{len(report_rows)} Submissions").font = bold_font
        for col_i in range(1, 17):
            ws.cell(row=cur_row, column=col_i).border = thin_border
            ws.cell(row=cur_row, column=col_i).fill = PatternFill(start_color="F1F5F9", end_color="F1F5F9", fill_type="solid")
            
        widths = {
            "A": 6, "B": 14, "C": 16, "D": 24, "E": 18, 
            "F": 28, "G": 42, "H": 22, "I": 22, "J": 22, 
            "K": 14, "L": 14, "M": 16, "N": 18, "O": 20, "P": 14
        }
        for col_letter, w in widths.items():
            ws.column_dimensions[col_letter].width = w
            
        stream = io.BytesIO()
        wb.save(stream)
        stream.seek(0)
        
        filename = f"GIMPA_Academic_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        return StreamingResponse(
            stream,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )



