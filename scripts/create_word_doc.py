import os
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_gimpa_documentation_docx(output_path):
    doc = Document()

    # Page setup - Margins (1 inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Palette Constants
    COLOR_PRIMARY = RGBColor(0x00, 0x33, 0x66)    # Deep GIMPA Blue (#003366)
    COLOR_SECONDARY = RGBColor(0xC5, 0x9B, 0x27)  # GIMPA Gold (#C59B27)
    COLOR_TEXT = RGBColor(0x2D, 0x37, 0x48)       # Charcoal (#2D3748)
    COLOR_MUTED = RGBColor(0x71, 0x80, 0x96)      # Muted Gray (#718096)
    
    HEX_PRIMARY = "003366"
    HEX_SECONDARY = "C59B27"
    HEX_BG_LIGHT = "F7FAFC"
    HEX_CALLOUT_BG = "F0F4F8"
    HEX_BORDER = "E2E8F0"

    # Base Normal Style setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = COLOR_TEXT

    # Helpers
    def set_cell_background(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
        tcPr.append(tcMar)

    def add_title(text, subtitle_text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(36)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(26)
        run.bold = True
        run.font.color.rgb = COLOR_PRIMARY

        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub.paragraph_format.space_before = Pt(0)
        p_sub.paragraph_format.space_after = Pt(24)
        run_sub = p_sub.add_run(subtitle_text)
        run_sub.font.name = 'Calibri'
        run_sub.font.size = Pt(14)
        run_sub.font.color.rgb = COLOR_SECONDARY
        run_sub.bold = True

        # Decorative line
        p_line = doc.add_paragraph()
        p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_line.paragraph_format.space_after = Pt(24)
        run_line = p_line.add_run("____________________________________________________")
        run_line.font.color.rgb = COLOR_SECONDARY

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(18)
        run.bold = True
        run.font.color.rgb = COLOR_PRIMARY

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.bold = True
        run.font.color.rgb = COLOR_SECONDARY

    def add_h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = COLOR_PRIMARY

    def add_p(text, bold_prefix=None, space_after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_pre = p.add_run(bold_prefix)
            r_pre.bold = True
            r_pre.font.color.rgb = COLOR_PRIMARY
        r = p.add_run(text)
        r.font.color.rgb = COLOR_TEXT
        return p

    def add_bullet(text, bold_prefix=None, level=0):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
        if bold_prefix:
            r_pre = p.add_run(bold_prefix)
            r_pre.bold = True
            r_pre.font.color.rgb = COLOR_PRIMARY
        r = p.add_run(text)
        r.font.color.rgb = COLOR_TEXT
        return p

    def add_callout(title, text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        cell = tbl.cell(0, 0)
        set_cell_background(cell, HEX_CALLOUT_BG)
        
        tcPr = cell._tc.get_or_add_tcPr()
        borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="none"/><w:left w:val="single" w:sz="36" w:space="0" w:color="{HEX_PRIMARY}"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>')
        tcPr.append(borders)
        set_cell_margins(cell, top=120, bottom=120, left=180, right=180)

        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        if title:
            r_t = p.add_run(f"NOTE: {title}\n")
            r_t.bold = True
            r_t.font.color.rgb = COLOR_PRIMARY
            r_t.font.size = Pt(10.5)

        r_b = p.add_run(text)
        r_b.font.size = Pt(10)
        r_b.font.color.rgb = COLOR_TEXT
        
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def format_table(tbl, col_widths=None):
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        for i, row in enumerate(tbl.rows):
            trPr = row._tr.get_or_add_trPr()
            trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

            if i == 0:
                trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
                for cell in row.cells:
                    set_cell_background(cell, HEX_PRIMARY)
                    set_cell_margins(cell, top=120, bottom=120, left=140, right=140)
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        p.paragraph_format.space_before = Pt(2)
                        p.paragraph_format.space_after = Pt(2)
                        for r in p.runs:
                            r.font.bold = True
                            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            r.font.size = Pt(10)
            else:
                bg = HEX_BG_LIGHT if i % 2 == 1 else "FFFFFF"
                for cell in row.cells:
                    set_cell_background(cell, bg)
                    set_cell_margins(cell, top=90, bottom=90, left=140, right=140)
                    tcPr = cell._tc.get_or_add_tcPr()
                    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/><w:right w:val="none"/></w:tcBorders>')
                    tcPr.append(borders)
                    for p in cell.paragraphs:
                        p.paragraph_format.space_before = Pt(2)
                        p.paragraph_format.space_after = Pt(2)
                        for r in p.runs:
                            r.font.size = Pt(9.5)
                            r.font.color.rgb = COLOR_TEXT

            if col_widths:
                for idx, w in enumerate(col_widths):
                    if idx < len(row.cells):
                        row.cells[idx].width = Inches(w)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # -------------------------------------------------------------
    # DOCUMENT CONTENT GENERATION
    # -------------------------------------------------------------

    # Title Banner
    add_title("GIMPA THESIS REPOSITORY SYSTEM", "Comprehensive Technical Manual & End-User System Guide")

    # Document Meta Info Box
    tbl_meta = doc.add_table(rows=4, cols=2)
    meta_data = [
        ("Institution", "Ghana Institute of Management and Public Administration (GIMPA)"),
        ("System Target", "Undergraduate & Postgraduate Thesis/Dissertation Management"),
        ("Document Version", "v1.0 (Production Documentation)"),
        ("Target Audience", "Students, Supervisors, HODs, Deans, Coordinators, Examiners, Librarians & System Administrators")
    ]
    for idx, (label, val) in enumerate(meta_data):
        row = tbl_meta.rows[idx]
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(val)
    format_table(tbl_meta, col_widths=[2.0, 4.5])

    doc.add_page_break()

    # SECTION 1: EXECUTIVE SUMMARY
    add_h1("1. Executive Summary & System Overview")
    add_p("The GIMPA Thesis Repository System is a comprehensive, enterprise-grade digital platform designed specifically for the Ghana Institute of Management and Public Administration (GIMPA). It automates and streamlines the complete thesis and dissertation research lifecycle across all undergraduate and postgraduate academic programs.")
    add_p("Historically, thesis administration at higher education institutions suffered from fragmented paper submissions, manual email handoffs, lack of progress visibility, lost feedback, and complex multi-examiner mark reconciliation. This system solves these operational challenges by consolidating all thesis activities into a unified, 5-phase structured workflow supported by Role-Based Access Control (RBAC), automated notifications, inline document review capabilities, and an institutional digital repository.")
    
    add_callout("Key Strategic Objective", "To replace email-based and manual paper-driven thesis reviews with an auditable, transparent, step-by-step electronic workflow that enforces institutional quality standards and protects institutional intellectual property.")

    # SECTION 2: SYSTEM ARCHITECTURE & TECH STACK
    add_h1("2. System Architecture & Technology Stack")
    add_p("The system is engineered following modern decoupled web architecture principles, featuring a high-performance RESTful API backend and a responsive single-page application (SPA) frontend.")

    add_h2("2.1 Technical Stack Breakdown")
    tbl_stack = doc.add_table(rows=7, cols=3)
    tbl_stack.rows[0].cells[0].paragraphs[0].add_run("Layer")
    tbl_stack.rows[0].cells[1].paragraphs[0].add_run("Technology / Library")
    tbl_stack.rows[0].cells[2].paragraphs[0].add_run("Architecture Rationale & Role")

    stack_rows = [
        ("Backend Framework", "FastAPI (Python 3.12+)", "High performance, asynchronous I/O, auto-generated Swagger/OpenAPI documentation, strict typing with Pydantic v2."),
        ("Database Engine", "SQLite / PostgreSQL with SQLAlchemy 2.0 ORM", "Relational schema ensuring data integrity across users, phase histories, examiner marks, and system audit logs."),
        ("Security & Auth", "OAuth2, JWT Tokens, Passlib (bcrypt)", "Stateless authentication with token refresh, embedding user role claims and department scope."),
        ("Frontend Application", "React 19 + React Router v7 + Vite", "Fast rendering engine with role-based component dashboards, persistent tab navigation, and zero-reload state updates."),
        ("Styling & Components", "Tailwind CSS v4 + Lucide Icons", "Custom institutional color tokens (Navy/Gold), accessible UI primitives, and responsive layout grids."),
        ("Document Engine", "python-docx, PyPDF2, ONLYOFFICE React", "Parsing and rendering student DOCX/PDF submissions, enabling inline browser viewing and automated extraction.")
    ]

    for idx, row in enumerate(stack_rows, start=1):
        tbl_stack.rows[idx].cells[0].paragraphs[0].add_run(row[0])
        tbl_stack.rows[idx].cells[1].paragraphs[0].add_run(row[1])
        tbl_stack.rows[idx].cells[2].paragraphs[0].add_run(row[2])
    
    format_table(tbl_stack, col_widths=[1.5, 2.2, 2.8])

    # SECTION 3: USER ROLES & PERMISSIONS MATRIX
    add_h1("3. User Roles & Permission Matrix")
    add_p("The system implements a strict Role-Based Access Control (RBAC) model across 8 distinct user roles. Each user role gets a custom dashboard tailored to their operational responsibilities.")

    tbl_rbac = doc.add_table(rows=9, cols=3)
    tbl_rbac.rows[0].cells[0].paragraphs[0].add_run("User Role")
    tbl_rbac.rows[0].cells[1].paragraphs[0].add_run("Scope of Access")
    tbl_rbac.rows[0].cells[2].paragraphs[0].add_run("Primary System Capabilities")

    rbac_rows = [
        ("Student", "Individual Thesis Record", "Submit thesis proposal, upload chapter steps (Steps 1-5), view supervisor feedback, upload final combined thesis draft and post-exam corrections."),
        ("Supervisor", "Assigned Students", "Approve/reject proposals, grade chapter step drafts, provide inline doc feedback & external remarks, approve post-examination corrections."),
        ("HOD (Head of Department)", "Departmental Theses", "Approve/reject proposals, assign supervisors, review examiner scores, assign internal/external examiners, grant final departmental sign-off."),
        ("Dean", "School-wide (Multi-Dept)", "Read-only analytics dashboard displaying research pipeline progress, supervisor workload, and completion bottleneck metrics across departments."),
        ("Project Coordinator", "Department/School", "Oversee examiner assignment, compile examiner scripts into ZIP archives, forward consolidated examiner remarks to students."),
        ("Internal Examiner", "Assigned Examination ZIPs", "Download compiled student work packages, upload Excel mark sheets and detailed assessment reports."),
        ("External Examiner", "Assigned Examination ZIPs", "Download assigned student packages, grade thesis against external benchmarks, submit external evaluation mark sheets."),
        ("Librarian / Public", "Global / Published Archive", "Review Phase 5 approved theses in publishing queue, publish theses to public repository, manage metadata, search public archive.")
    ]

    for idx, row in enumerate(rbac_rows, start=1):
        tbl_rbac.rows[idx].cells[0].paragraphs[0].add_run(row[0])
        tbl_rbac.rows[idx].cells[1].paragraphs[0].add_run(row[1])
        tbl_rbac.rows[idx].cells[2].paragraphs[0].add_run(row[2])

    format_table(tbl_rbac, col_widths=[1.5, 1.8, 3.2])

    # SECTION 4: THE 5-PHASE THESIS WORKFLOW ENGINE
    add_h1("4. Multi-Phase Thesis Workflow Engine")
    add_p("The core of the system is a 5-phase sequential state machine that enforces strict academic rigor and institutional oversight at every stage of research.")

    add_h2("Phase 1: Topic & Proposal Submission")
    add_bullet("Student submits topic title, research abstract, department selection, and requested supervisor.", "Step 1.1: Proposal Creation — ")
    add_bullet("Assigned supervisor evaluates proposal feasibility and either approves or requests revisions.", "Step 1.2: Supervisor Review — ")
    add_bullet("Department HOD grants formal approval and officially confirms supervisor assignment.", "Step 1.3: HOD Approval — ")

    add_h2("Phase 2: Chapter-by-Chapter Review (Steps 1 to 5)")
    add_p("Phase 2 replaces manual chapter submissions with a structured 5-step draft progression:")
    add_bullet("Introduction, Problem Statement & Research Objectives", "Step 1: Chapter 1 — ")
    add_bullet("Literature Review & Theoretical Framework", "Step 2: Chapter 2 — ")
    add_bullet("Research Methodology, Data Collection & Design", "Step 3: Chapter 3 — ")
    add_bullet("Data Analysis, Presentation & Findings", "Step 4: Chapter 4 — ")
    add_bullet("Discussion, Conclusion & Recommendations", "Step 5: Chapter 5 — ")
    add_p("For each step, the student uploads a document draft. The supervisor reviews the submission online and issues a decision: Approved or Revise. Resubmissions preserve revision history without overwriting previous attempts. Once all 5 steps are approved, the supervisor clicks 'Finish Steps' to advance the project to Phase 3.")

    add_h2("Phase 3: Examination & Examiner ZIP Compilation")
    add_bullet("Student uploads the full, combined final thesis document.", "Draft Submission — ")
    add_bullet("HOD or Coordinator assigns minimum 1 Internal Examiner and 1 External Examiner.", "Examiner Assignment — ")
    add_bullet("System automatically generates a downloadable ZIP file containing assigned thesis drafts for examiners.", "ZIP Package Build — ")
    add_bullet("Examiners assess work and upload Excel mark sheets + qualitative feedback.", "Script Evaluation — ")
    add_bullet("Coordinator reviews and forwards consolidated examiner remarks to student.", "Feedback Relay — ")

    add_h2("Phase 4: Post-Examination Corrections")
    add_bullet("Student addresses examiner feedback and uploads revised thesis corrections script.", "Corrections Upload — ")
    add_bullet("Requires dual sign-off: Supervisor approval followed by HOD final sign-off.", "Dual Approval — ")

    add_h2("Phase 5: Publishing & Public Archiving")
    add_bullet("Approved thesis is queued for Library cataloging.", "Library Queue — ")
    add_bullet("Librarian verifies metadata and marks thesis as 'Published'.", "Public Archive — ")
    add_bullet("Thesis becomes searchable in the GIMPA Public Repository with online preview and citation export.", "Repository Access — ")

    # SECTION 5: ROLE-BASED USER GUIDES
    add_h1("5. Comprehensive Role-Based User Manuals")

    add_h2("5.1 Student User Manual")
    add_p("As a student, the platform guides you step-by-step from proposal creation to graduation archiving:")
    add_bullet("Log in using your student credentials (email/password).", "1. Accessing Portal: ")
    add_bullet("Click 'Submit Proposal', fill in topic details, select your academic department, and choose a supervisor.", "2. Creating Proposal: ")
    add_bullet("Navigate to your active thesis dashboard. Under Phase 2, click 'Upload Step File' for Chapter 1.", "3. Submitting Chapter Steps: ")
    add_bullet("If your supervisor requests revisions, click 'Submit Edited Step File' on the specific step card to submit updated drafts.", "4. Handling Revisions: ")
    add_bullet("Once all 5 steps are marked 'Approved', upload your combined thesis under Phase 3 for examination.", "5. Final Compilation: ")
    add_bullet("Download examiner feedback in Phase 4, make requested edits, and submit your final correction file for graduation sign-off.", "6. Post-Exam Corrections: ")

    add_h2("5.2 Supervisor User Manual")
    add_p("Supervisors manage student progress and provide actionable feedback directly through the browser:")
    add_bullet("View all assigned students on your Supervisor Dashboard with current phase status and pending tasks.", "1. Dashboard Overview: ")
    add_bullet("Review student proposals in Phase 1 and mark them 'Approved' or 'Revise'.", "2. Proposal Evaluation: ")
    add_bullet("Click on any submitted Chapter Step (Steps 1-5) to inspect uploaded DOCX/PDF files.", "3. Step Inspection: ")
    add_bullet("Use the Online Document Viewer to view student work. Enter detailed guidance in the 'External Supervisor Remarks' box.", "4. Giving Feedback: ")
    add_bullet("When all 5 chapter steps are completed satisfactorily, click the 'Finish Steps' button to authorize Phase 3 examination.", "5. Finishing Steps: ")

    add_h2("5.3 Head of Department (HOD) User Manual")
    add_p("HODs maintain academic quality standards and manage departmental research allocations:")
    add_bullet("Review incoming student proposals for department alignment.", "1. Proposal Approval: ")
    add_bullet("Assign or reassign supervisors based on topic domain and workload capacity.", "2. Supervisor Allocation: ")
    add_bullet("Approve internal and external examiner assignments made by coordinators.", "3. Examiner Oversight: ")
    add_bullet("Review multi-examiner mark reconciliations and give final departmental approval in Phase 4.", "4. Final Departmental Sign-off: ")

    add_h2("5.4 Project Coordinator & Admin User Manual")
    add_p("Coordinators manage the operational logistics of the examination workflow:")
    add_bullet("Assign internal and external examiners to eligible Phase 3 thesis submissions.", "1. Assigning Examiners: ")
    add_bullet("Click 'Generate Examiner ZIP' to automatically bundle thesis files for external assessors.", "2. Generating Assessment ZIPs: ")
    add_bullet("Upload returned examiner scorecards (Excel) and narrative scripts.", "3. Processing Marksheets: ")
    add_bullet("Consolidate feedback items and release them to student dashboards.", "4. Relaying Feedback: ")

    add_h2("5.5 Examiner User Manual (Internal & External)")
    add_bullet("Log into the secure Examiner Portal.", "1. Secure Login: ")
    add_bullet("Download the generated ZIP archive containing assigned student thesis documents.", "2. Downloading ZIP Package: ")
    add_bullet("Review manuscripts offline or online.", "3. Manuscript Assessment: ")
    add_bullet("Upload completed Excel evaluation forms and summary comments for each student.", "4. Submitting Grades: ")

    add_h2("5.6 Librarian User Manual")
    add_bullet("Access the Phase 5 Library Publishing Queue.", "1. Publishing Queue: ")
    add_bullet("Verify abstract, author details, department taxonomy, publication year, and keywords.", "2. Metadata Verification: ")
    add_bullet("Click 'Publish Thesis' to render the paper live in the public archive.", "3. Publishing Action: ")

    # SECTION 6: API REFERENCE & ENDPOINTS
    add_h1("6. REST API Reference & Endpoint Specifications")
    add_p("The backend exposes a secure REST API formatted under `/api/v1`. Key endpoints include:")

    tbl_api = doc.add_table(rows=10, cols=3)
    tbl_api.rows[0].cells[0].paragraphs[0].add_run("HTTP Method & Endpoint")
    tbl_api.rows[0].cells[1].paragraphs[0].add_run("Auth & Role Scope")
    tbl_api.rows[0].cells[2].paragraphs[0].add_run("Endpoint Function & Behavior")

    api_endpoints = [
        ("POST /api/v1/auth/login", "Public", "Authenticate user credentials, issue JWT access token and refresh token."),
        ("GET /api/v1/auth/me", "Authenticated", "Return current authenticated user profile, role claims, and department id."),
        ("POST /api/v1/theses/submit-proposal", "Student", "Create a new thesis entry and submit initial Phase 1 proposal."),
        ("POST /api/v1/theses/{id}/submit-step", "Student", "Upload new chapter step draft file (Steps 1 to 5)."),
        ("POST /api/v1/theses/steps/{step_id}/resubmit", "Student", "Resubmit revised step file while retaining step history."),
        ("POST /api/v1/steps/{step_id}/decision", "Supervisor", "Submit supervisor decision (approved / revise) and comments."),
        ("POST /api/v1/theses/{id}/finish-steps", "Supervisor", "Complete Phase 2 chapter steps and advance project to Phase 3."),
        ("GET /api/v1/theses/steps/{step_id}/file", "Auth (Role-scoped)", "Securely stream and download submitted DOCX or PDF file."),
        ("POST /api/v1/users/import", "Admin / Coordinator", "Bulk import user accounts from CSV or Excel roster files.")
    ]

    for idx, row in enumerate(api_endpoints, start=1):
        tbl_api.rows[idx].cells[0].paragraphs[0].add_run(row[0])
        tbl_api.rows[idx].cells[1].paragraphs[0].add_run(row[1])
        tbl_api.rows[idx].cells[2].paragraphs[0].add_run(row[2])

    format_table(tbl_api, col_widths=[2.2, 1.5, 2.8])

    # SECTION 7: SYSTEM SETUP, DEPLOYMENT & MAINTENANCE
    add_h1("7. System Setup & Operations Guide")
    add_p("Follow these instructions to install, configure, and run the GIMPA Thesis Repository System locally or on a server.")

    add_h2("7.1 Backend Setup Instructions")
    add_p("1. Navigate to backend directory, create Python virtual environment, and install dependencies:")
    add_p("cd backend\npython -m venv venv\n.\\venv\\Scripts\\activate   # On Windows\nsource venv/bin/activate  # On Linux/macOS\npip install -r requirements.txt", bold_prefix="Command Sequence:\n")

    add_p("2. Configure environment variables in `.env`:")
    add_p("DATABASE_URL=sqlite:///./sql_app.db\nJWT_SECRET=your_super_secret_key_here\nSMTP_ENABLED=true\nSMTP_HOST=smtp.gimpa.edu.gh\nSMTP_PORT=587\nSMTP_USERNAME=notifications@gimpa.edu.gh\nSMTP_PASSWORD=your_smtp_password", bold_prefix="Environment Config:\n")

    add_p("3. Initialize database tables and start FastAPI uvicorn server:")
    add_p("python -m app.db.init_db\nuvicorn app.main:app --reload --host 127.0.0.1 --port 8011", bold_prefix="Launch Server:\n")

    add_h2("7.2 Frontend Setup Instructions")
    add_p("1. Navigate to frontend directory and install dependencies:")
    add_p("cd frontend\nnpm install\nnpm run dev", bold_prefix="Launch Frontend:\n")
    add_p("The web app will run on http://localhost:5173 and automatically proxy requests to the backend server on port 8011.")

    add_h2("7.3 Seed Data & Demo Accounts")
    add_p("To test the complete workflow out of the box, run the included database seeding scripts:")
    add_bullet("Populates GIMPA Schools and Academic Departments.", "python scripts/seed_departments.py — ")
    add_bullet("Creates test accounts for Student, Supervisor, HOD, Dean, Coordinator, Examiner, and Librarian.", "python scripts/seed_spec_accounts.py — ")
    add_bullet("Seeds mock thesis proposals at every phase for workflow testing.", "python scripts/seed_workflow_demo.py — ")

    # SECTION 8: CONCLUSION & REPOSITORY METADATA
    add_h1("8. Document Control & Repository Info")
    add_p("This document represents the official system manual for the GIMPA Thesis Repository System. All updates to the codebase or operational workflows should be documented in subsequent revisions of this manual.")
    
    add_callout("Repository Location", "GitHub Repository URL: https://github.com/Lexies99/GIMPA_Thesis_Repository.git\nMaintained by: GIMPA Thesis System Engineering Team")

    # Save Document
    doc.save(output_path)
    print(f"Document successfully created at: {output_path}")

if __name__ == "__main__":
    output_filename = "GIMPA_Thesis_Repository_System_User_Documentation.docx"
    target_path = os.path.join(r"d:\NSS\GIMPA_Thesis_Repository_Bootstrap", output_filename)
    create_gimpa_documentation_docx(target_path)
